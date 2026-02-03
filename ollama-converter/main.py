"""
Ollama-based Data Converter API
Converts unstructured data (text, CSV, JSON, etc.) into structured Startup/Investor format
"""

from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any, Tuple, AsyncGenerator
import ollama
import os
import httpx
import json
import re
from io import StringIO
import csv
import asyncio
from typing import Tuple, List as TypingList

app = FastAPI(title="Ollama Data Converter API")

# Limit how much extracted text we send to the model (large PDFs often cause truncated JSON output).
# This keeps responses short enough to remain valid JSON.
MAX_MODEL_INPUT_CHARS = int(os.environ.get("MAX_MODEL_INPUT_CHARS", "24000"))

# OCR settings (for scanned/image PDFs)
OCR_MAX_PAGES = int(os.environ.get("OCR_MAX_PAGES", "5"))
OCR_DPI = int(os.environ.get("OCR_DPI", "200"))
# Limit PDF pages to reduce timeouts on large files
MAX_PDF_PAGES = int(os.environ.get("MAX_PDF_PAGES", "8"))
# Claude Vision API for complex PDFs (uses existing ANTHROPIC_API_KEY)
# Parallel processing settings
MAX_PARALLEL_PAGES = int(os.environ.get("MAX_PARALLEL_PAGES", "10"))


async def extract_with_claude_vision(page_images: List[bytes]) -> str:
    """
    Use Claude 3.5 Sonnet Vision to extract text from PDF page images.
    Best for: Scanned PDFs, complex layouts, tables with merged cells.
    Uses existing ANTHROPIC_API_KEY.
    """
    if not ANTHROPIC_API_KEY:
        return None
    
    import base64
    
    parts = []
    for idx, img_bytes in enumerate(page_images, start=1):
        try:
            # Convert image bytes to base64
            img_base64 = base64.b64encode(img_bytes).decode('utf-8')
            
            prompt = """Extract all text from this document page. Preserve:
- Table structure (use markdown tables if possible)
- Headers and footers
- Multi-column layouts
- Bullet points and numbered lists
- All numbers, dates, and proper nouns

Return the extracted text in a clear, structured format."""
            
            # Use Claude Vision API
            headers = {
                "Content-Type": "application/json",
                "x-api-key": ANTHROPIC_API_KEY,
                "anthropic-version": "2023-06-01",
            }
            url = get_anthropic_api_url()
            
            payload = {
                "model": "claude-3-5-sonnet-20240620",  # Claude 3.5 Sonnet with vision
                "max_tokens": 4096,
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image",
                                "source": {
                                    "type": "base64",
                                    "media_type": "image/png",
                                    "data": img_base64
                                }
                            },
                            {
                                "type": "text",
                                "text": prompt
                            }
                        ]
                    }
                ]
            }
            
            async with httpx.AsyncClient(timeout=60.0) as client:
                res = await client.post(url, headers=headers, json=payload)
                res.raise_for_status()
                data = res.json()
                content = data.get("content", [])
                if isinstance(content, list) and content:
                    page_text = content[0].get("text", "") if content[0].get("type") == "text" else ""
                else:
                    page_text = str(content) if content else ""
                
                if page_text:
                    parts.append(f"\n--- Page {idx} (Claude Vision) ---\n{page_text}")
        except Exception as e:
            print(f"Claude Vision extraction failed for page {idx}: {e}")
            parts.append(f"\n--- Page {idx} (Claude Vision failed: {e}) ---\n")
    
    return "\n".join(parts).strip() if parts else None


def try_ocr_pdf_bytes(content: bytes) -> str:
    """
    Best-effort OCR for scanned/image-only PDFs.
    Requires:
      - pytesseract (python)
      - pdf2image (python)
      - Tesseract installed on the OS
      - Poppler installed on the OS (Windows: poppler-utils)
    """
    try:
        from pdf2image import convert_from_bytes  # type: ignore
        import pytesseract  # type: ignore
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=(
                "Scanned/image PDF detected but OCR dependencies are missing. "
                "Install: pip install pytesseract pdf2image, then install Tesseract + Poppler on Windows. "
                f"Error: {str(e)}"
            ),
        )

    try:
        images = convert_from_bytes(content, dpi=OCR_DPI, first_page=1, last_page=max(1, OCR_MAX_PAGES))
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=(
                "Failed to render PDF pages for OCR. On Windows you usually need Poppler installed and on PATH. "
                f"Error: {str(e)}"
            ),
        )

    parts: List[str] = []
    for idx, img in enumerate(images, start=1):
        try:
            txt = pytesseract.image_to_string(img) or ""
            txt = re.sub(r"\s+\n", "\n", txt)
            txt = re.sub(r"\n{3,}", "\n\n", txt).strip()
            parts.append(f"\n--- OCR Page {idx} ---\n{txt}")
        except Exception as e:
            parts.append(f"\n--- OCR Page {idx} ---\n[OCR_FAILED: {str(e)}]")

    return "\n".join(parts).strip()

# Converter provider settings
_provider_env = os.getenv("CONVERTER_PROVIDER")
CONVERTER_PROVIDER = (_provider_env or "ollama").lower().strip()

# Ollama connection settings
OLLAMA_HOST = os.getenv("OLLAMA_HOST", "http://127.0.0.1:11434").rstrip("/")
PREFERRED_OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "vc-converter:latest")

# Anthropic (Claude) settings
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY")
ANTHROPIC_MODEL = os.getenv("ANTHROPIC_MODEL", "claude-3-5-sonnet-20240620")
ANTHROPIC_API_URL = os.getenv("ANTHROPIC_API_URL", "https://api.anthropic.com/v1/messages")

# Known invalid models to exclude
INVALID_MODELS = {"claude-3-5-sonnet-20241022"}

# Filter out invalid models from env var if set
if ANTHROPIC_MODEL in INVALID_MODELS:
    ANTHROPIC_MODEL = "claude-3-5-sonnet-20240620"

ANTHROPIC_MODEL_FALLBACKS = [
    m for m in [
        ANTHROPIC_MODEL,
        "claude-3-5-sonnet-20240620",
        "claude-3-5-haiku-20241022",
        "claude-3-haiku-20240307",
        "claude-3-opus-20240229",
    ] if m not in INVALID_MODELS
]

# Ask-the-fund settings (generous tokens for comprehensive answers)
ASK_MAX_TOKENS = int(os.getenv("ASK_MAX_TOKENS", "4000"))  # Increased from 1000 for more detailed responses
ASK_MAX_SOURCES = int(os.getenv("ASK_MAX_SOURCES", "5"))   # More sources for better context
ASK_MAX_SNIPPET_CHARS = int(os.getenv("ASK_MAX_SNIPPET_CHARS", "500"))  # Larger snippets for better answers
# Use Haiku for simple questions (3-5x faster, 75% cheaper)
USE_HAIKU_FOR_SIMPLE = os.getenv("USE_HAIKU_FOR_SIMPLE", "true").lower() == "true"

# Embeddings settings (semantic search)
EMBEDDINGS_PROVIDER = os.getenv("EMBEDDINGS_PROVIDER", "voyage").lower().strip()
OLLAMA_EMBEDDING_MODEL = os.getenv("OLLAMA_EMBEDDING_MODEL", "nomic-embed-text")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OPENAI_EMBEDDING_MODEL = os.getenv("OPENAI_EMBEDDING_MODEL", "text-embedding-3-small")  # 1536 dimensions
VOYAGE_API_KEY = os.getenv("VOYAGE_API_KEY")
VOYAGE_EMBEDDING_MODEL = os.getenv("VOYAGE_EMBEDDING_MODEL", "voyage-finance-2")
EMBEDDING_DIM = int(os.getenv("EMBEDDING_DIM", "1536"))

# Reranking settings (cross-encoder)
COHERE_API_KEY = os.getenv("COHERE_API_KEY")
RERANK_MODEL = os.getenv("RERANK_MODEL", "rerank-english-v3.0")

# Ingestion settings
CLICKUP_API_TOKEN = os.getenv("CLICKUP_API_TOKEN")

def get_anthropic_api_url() -> str:
    """
    Normalize Anthropic API URL.
    Accepts base URLs like:
      - https://api.anthropic.com
      - https://api.anthropic.com/v1
      - https://api.anthropic.com/v1/messages
      - https://api.anthropic.com/v1/messages/
    Returns the full /v1/messages endpoint.
    """
    base = (ANTHROPIC_API_URL or "").strip()
    if not base:
        return "https://api.anthropic.com/v1/messages"
    base = base.rstrip("/")
    if base.endswith("/v1/messages"):
        return base
    if base.endswith("/v1"):
        return f"{base}/messages"
    return f"{base}/v1/messages"


async def fetch_ollama_model_names() -> List[str]:
    """
    More reliable than python ollama.list() on some setups.
    Uses Ollama's HTTP API to list installed models.
    """
    names: List[str] = []

    # First, try the HTTP /api/tags endpoint
    try:
        async with httpx.AsyncClient(timeout=2.0) as client:
            res = await client.get(f"{OLLAMA_HOST}/api/tags")
            res.raise_for_status()
            data = res.json() or {} 
            models = data.get("models", []) or []
            for m in models:
                if isinstance(m, dict) and m.get("name"):
                    names.append(m["name"])
    except Exception:
        # swallow and try python client fallback below
        names = []

    # Fallback: python client list() if HTTP returned nothing
    if not names:
        try:
            client = get_ollama_client()
            models = client.list()
            if isinstance(models, dict):
                for m in models.get("models", []) or []:
                    if isinstance(m, dict) and m.get("name"):
                        names.append(m["name"])
                    elif isinstance(m, str):
                        names.append(m)
        except Exception:
            pass

    return names


def pick_model(available_models: List[str]) -> str:
    """
    Pick a model name to use for conversion.
    Prefer env OLLAMA_MODEL, then vc-converter*, then llama3.1*, then llama3.2*, else first.
    """
    if not available_models:
        return PREFERRED_OLLAMA_MODEL

    if PREFERRED_OLLAMA_MODEL in available_models:
        return PREFERRED_OLLAMA_MODEL

    for prefix in ["vc-converter", "llama3.1", "llama3.2", "llama3"]:
        for name in available_models:
            if name.startswith(prefix):
                return name

    return available_models[0]


def get_ollama_client() -> "ollama.Client":
    # Force the host so the python client matches what `ollama list` uses.
    return ollama.Client(host=OLLAMA_HOST)

# CORS middleware to allow frontend requests
# FastAPI CORSMiddleware doesn't support ["*"] with allow_credentials=False
# So we use a list of common origins or allow all by using None
_cors_origins_env = os.getenv("CORS_ALLOW_ORIGINS", "*")
if _cors_origins_env.strip() == "*":
    # Use None to allow all origins (FastAPI's way)
    cors_allow_origins = None
else:
    cors_allow_origins = [o.strip() for o in _cors_origins_env.split(",") if o.strip()]

app.add_middleware(
    CORSMiddleware,
    # None means allow all origins, otherwise use the list
    allow_origins=cors_allow_origins if cors_allow_origins else ["*"],
    allow_credentials=False,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["*"],
    expose_headers=["*"],
)

# Data models
class StartupData(BaseModel):
    companyName: str
    geoMarkets: List[str]
    industry: str
    fundingTarget: int
    fundingStage: str
    availabilityStatus: str = "present"

class InvestorData(BaseModel):
    firmName: str
    memberName: str
    geoFocus: List[str]
    industryPreferences: List[str]
    stagePreferences: List[str]
    minTicketSize: int
    maxTicketSize: int
    totalSlots: int = 3
    tableNumber: Optional[str] = None
    availabilityStatus: str = "present"

class MentorData(BaseModel):
    fullName: str
    email: str
    linkedinUrl: Optional[str] = None
    geoFocus: List[str]
    industryPreferences: List[str]
    expertiseAreas: List[str]
    totalSlots: int = 3
    availabilityStatus: str = "present"

class CorporateData(BaseModel):
    firmName: str
    contactName: str
    email: Optional[str] = None
    geoFocus: List[str]
    industryPreferences: List[str]
    partnershipTypes: List[str]
    stages: List[str]
    totalSlots: int = 3
    availabilityStatus: str = "present"

class ConversionRequest(BaseModel):
    data: str  # Unstructured data (text, CSV, JSON, etc.)
    dataType: Optional[str] = None  # 'startup', 'investor', or None for auto-detect
    format: Optional[str] = None  # 'csv', 'text', 'json', etc.

class ConversionResponse(BaseModel):
    startups: List[StartupData] = []
    investors: List[InvestorData] = []
    mentors: List[MentorData] = []
    corporates: List[CorporateData] = []
    detectedType: str
    confidence: float
    warnings: List[str] = []
    errors: List[str] = []
    raw_content: Optional[str] = None

class FileValidationResponse(BaseModel):
    isValid: bool
    errors: List[str] = []
    warnings: List[str] = []
    detectedType: Optional[str] = None
    startupCsvTemplate: Optional[str] = None
    investorCsvTemplate: Optional[str] = None

class ClickUpIngestRequest(BaseModel):
    list_id: str
    include_closed: bool = True

class ClickUpIngestResponse(BaseModel):
    tasks: List[Dict[str, Any]] = []

class ClickUpListsRequest(BaseModel):
    team_id: str

class AskSource(BaseModel):
    title: Optional[str] = None
    snippet: Optional[str] = None
    file_name: Optional[str] = None

class AskDecision(BaseModel):
    startup_name: Optional[str] = None
    action_type: Optional[str] = None
    outcome: Optional[str] = None
    notes: Optional[str] = None

class ChatMessage(BaseModel):
    role: str  # "user" or "assistant"
    content: str

class AskRequest(BaseModel):
    question: str
    sources: List[AskSource] = []
    decisions: List[AskDecision] = []
    previous_messages: List[ChatMessage] = Field(default_factory=list, alias="previousMessages")

    class Config:
        allow_population_by_field_name = True

class AskResponse(BaseModel):
    answer: str

class RerankDocument(BaseModel):
    id: str
    text: str

class RerankRequest(BaseModel):
    query: str
    documents: List[RerankDocument]
    top_n: int | None = None

class RerankResult(BaseModel):
    id: str
    score: float

class RerankResponse(BaseModel):
    results: List[RerankResult]

class EmbedRequest(BaseModel):
    text: str
    input_type: Optional[str] = None

class EmbedResponse(BaseModel):
    embedding: List[float]

class RewriteQueryRequest(BaseModel):
    question: str
    previous_messages: Optional[List[ChatMessage]] = []

class RewriteQueryResponse(BaseModel):
    rewritten_question: str

class ClickUpListsResponse(BaseModel):
    lists: List[Dict[str, Any]] = []

class GoogleDriveIngestRequest(BaseModel):
    url: str
    access_token: Optional[str] = None

class GoogleDriveIngestResponse(BaseModel):
    title: str
    content: str
    raw_content: str  # Alias for content, for clarity
    sourceType: str

# System prompt for Ollama
SYSTEM_PROMPT = """You are a data extraction and conversion expert. Your task is to extract structured information from unstructured text and convert it into JSON format.

You will receive unstructured data about startups, investors, mentors, or corporates, and you must extract the following information:

FOR STARTUPS:
- companyName: The name of the company/startup
- geoMarkets: List of geographic markets (e.g., ["North America", "Europe"])
- industry: Industry sector (e.g., "AI/ML", "Fintech", "Healthtech")
- fundingTarget: Funding amount as integer (remove currency symbols, commas)
- fundingStage: Stage (e.g., "Pre-seed", "Seed", "Series A", "Series B+")

FOR INVESTORS:
- firmName: Name of the VC firm/investor
- memberName: The specific investor team member/person name (REQUIRED)
- geoFocus: List of geographic focus areas
- industryPreferences: List of preferred industries
- stagePreferences: List of preferred funding stages
- minTicketSize: Minimum investment amount as integer
- maxTicketSize: Maximum investment amount as integer
- totalSlots: Number of meeting slots (default: 3)
- tableNumber: Optional table/booth number

FOR MENTORS:
- fullName: Full name of the mentor (REQUIRED)
- email: Email address (REQUIRED)
- linkedinUrl: LinkedIn profile URL
- geoFocus: List of geographic focus areas
- industryPreferences: List of preferred industries
- expertiseAreas: List of expertise areas (e.g., ["Product Development", "Fundraising"])
- totalSlots: Number of meeting slots (default: 3)

FOR CORPORATES:
- firmName: Name of the corporate/company (REQUIRED)
- contactName: Name of the corporate contact person (REQUIRED)
- email: Email address
- geoFocus: List of geographic focus areas
- industryPreferences: List of preferred industries
- partnershipTypes: List of partnership types (e.g., ["Pilot Program", "Distribution"])
- stages: List of startup stages of interest (e.g., ["Seed", "Series A"])
- totalSlots: Number of meeting slots (default: 3)

IMPORTANT RULES:
1. Always return valid JSON only, no markdown or explanations
2. If multiple entities are found, return an array
3. Extract numbers from text (e.g., "$2M" -> 2000000, "€500K" -> 500000)
4. Parse lists from text (e.g., "North America, Europe" -> ["North America", "Europe"])
5. If information is missing, use reasonable defaults or empty arrays
6. For funding stages, normalize to: "Pre-seed", "Seed", "Series A", "Series B+"
7. For industries, use standard names: "Fintech", "Healthtech", "EdTech", "E-commerce", "Construction", "Transportation/Mobility", "AI/ML", "Logistics", "Consumer Goods", "SaaS", "CleanTech"

Return ONLY the JSON object or array, nothing else."""

def create_conversion_prompt(data: str, data_type: Optional[str] = None) -> str:
    """Create a prompt for Ollama to convert unstructured data"""
    # Keep prompt/model input bounded to reduce truncation.
    trimmed = data if len(data) <= MAX_MODEL_INPUT_CHARS else data[:MAX_MODEL_INPUT_CHARS]
    if len(trimmed) != len(data):
        trimmed = trimmed + "\n\n[TRUNCATED INPUT: content was longer than MAX_MODEL_INPUT_CHARS]"
    if data_type:
        prompt = (
            f"Extract {data_type} information from the following data and convert to JSON format.\n"
            f"Return ONLY valid JSON (no commentary). Keep it minimal: only the required fields.\n\n{trimmed}\n"
        )
    else:
        prompt = (
            "Extract startup or investor information from the following data. Auto-detect the type and convert to JSON.\n"
            "Return ONLY valid JSON (no commentary). Keep it minimal: only the required fields.\n\n"
            f"{trimmed}\n"
        )
    return prompt

def parse_ollama_response(response: str) -> Dict[str, Any]:
    """Parse model response and extract JSON"""
    # Remove markdown code blocks if present
    response = re.sub(r'```json\n?', '', response)
    response = re.sub(r'```\n?', '', response)
    response = response.strip()

    def extract_first_json_block(text: str) -> Optional[str]:
        """
        Extract the first complete JSON object/array from text using bracket balancing.
        This is robust against extra prose before/after JSON and avoids greedy regex traps.
        """
        start_idx = None
        stack: List[str] = []
        in_string = False
        escape = False

        for i, ch in enumerate(text):
            if start_idx is None:
                if ch == '{':
                    start_idx = i
                    stack = ['}']
                elif ch == '[':
                    start_idx = i
                    stack = [']']
                continue

            # We are inside a candidate JSON block
            if in_string:
                if escape:
                    escape = False
                elif ch == '\\':
                    escape = True
                elif ch == '"':
                    in_string = False
                continue

            if ch == '"':
                in_string = True
                continue

            if ch == '{':
                stack.append('}')
            elif ch == '[':
                stack.append(']')
            elif ch in ('}', ']'):
                if stack and ch == stack[-1]:
                    stack.pop()
                    if not stack and start_idx is not None:
                        return text[start_idx:i + 1]
                else:
                    # Mismatched closing bracket; keep scanning but this block is likely invalid.
                    pass

        return None

    # First try: extract a complete JSON block from within the response
    block = extract_first_json_block(response)
    if block:
        try:
            return json.loads(block)
        except json.JSONDecodeError:
            pass

    # Fallback: try parsing the whole response as JSON
    try:
        return json.loads(response)
    except json.JSONDecodeError:
        raise ValueError(
            "Could not parse JSON from model response (likely truncated / non-JSON). "
            f"Response starts with: {response[:200]}"
        )

def is_model_not_found(response: httpx.Response) -> bool:
    try:
        data = response.json() or {}
        err = data.get("error") or {}
        return err.get("type") == "not_found_error" and "model" in str(err.get("message", "")).lower()
    except Exception:
        return False

async def call_anthropic(prompt: str) -> str:
    """
    Call Anthropic (Claude) API and return the raw text response.
    Uses server-side API key from environment variables.
    """
    if not ANTHROPIC_API_KEY:
        raise HTTPException(
            status_code=503,
            detail="ANTHROPIC_API_KEY not set. Set it in the server environment to use Claude."
        )

    headers = {
        "Content-Type": "application/json",
        "x-api-key": ANTHROPIC_API_KEY,
        "anthropic-version": "2023-06-01",
    }
    url = get_anthropic_api_url()
    default_url = "https://api.anthropic.com/v1/messages"

    last_error: Optional[str] = None
    async with httpx.AsyncClient(timeout=60.0) as client:
        for model_name in [m for m in ANTHROPIC_MODEL_FALLBACKS if m]:
            payload = {
                "model": model_name,
                "max_tokens": 4096,
                "temperature": 0.1,
                "system": SYSTEM_PROMPT,
                "messages": [
                    {"role": "user", "content": prompt}
                ],
            }

            res = await client.post(url, headers=headers, json=payload)
            # If a misconfigured URL causes 404, retry with the canonical endpoint.
            if res.status_code == 404 and url != default_url:
                res = await client.post(default_url, headers=headers, json=payload)
            if res.status_code == 404 and is_model_not_found(res):
                last_error = f"Model not found: {model_name}"
                continue
            if res.status_code == 404:
                body = res.text[:400].strip()
                raise HTTPException(
                    status_code=502,
                    detail=(
                        f"Claude API 404 at {url}. Check ANTHROPIC_API_URL and account access. "
                        f"Response: {body or 'empty'}"
                    ),
                )
            try:
                res.raise_for_status()
            except httpx.HTTPError as e:
                last_error = str(e)
                continue

            data = res.json()
            content = data.get("content", [])
            if not content or not isinstance(content, list) or "text" not in content[0]:
                last_error = "Claude returned empty content."
                continue
            return content[0]["text"]

    raise HTTPException(
        status_code=502,
        detail=f"Claude API error: {last_error or 'Unknown error. All models failed.'}"
    )

def is_comprehensive_question(question: str) -> bool:
    """Detect if user wants a comprehensive answer (all you know, everything, detailed, etc.)"""
    q_lower = question.lower()
    comprehensive_patterns = [
        "all you know",
        "everything",
        "comprehensive",
        "detailed",
        "full",
        "complete",
        "tell me all",
        "tell me more all",
        "what do you know",
        "what can you tell me",
        "summarize",
        "overview",
        "what is inside",
        "what's inside",
        "what is in",
        "what's in",
        "what does it contain",
        "what does it say",
        "what is the content",
        "what are the contents",
        "just tell",
        "tell what",
        "all about",
        "everything about",
        "from these",
        "from the",
        "from source",
    ]
    return any(pattern in q_lower for pattern in comprehensive_patterns)

def is_raw_text_request(question: str) -> bool:
    """Detect if user is asking for raw/exact text from sources."""
    q_lower = question.lower()
    raw_patterns = [
        "raw text",
        "exact text",
        "verbatim",
        "just text",
        "full text",
        "show the text",
        "give me the text",
        "original text",
        "word for word",
    ]
    return any(pattern in q_lower for pattern in raw_patterns)


def levenshtein_distance(s1: str, s2: str) -> int:
    """
    Calculate Levenshtein distance between two strings.
    Used for fuzzy name matching.
    """
    if len(s1) < len(s2):
        return levenshtein_distance(s2, s1)
    if len(s2) == 0:
        return len(s1)
    
    previous_row = range(len(s2) + 1)
    for i, c1 in enumerate(s1):
        current_row = [i + 1]
        for j, c2 in enumerate(s2):
            insertions = previous_row[j + 1] + 1
            deletions = current_row[j] + 1
            substitutions = previous_row[j] + (c1 != c2)
            current_row.append(min(insertions, deletions, substitutions))
        previous_row = current_row
    
    return previous_row[-1]


def fuzzy_match_name(query_name: str, document_name: str, max_distance: int = 2) -> bool:
    """
    Check if two names match with fuzzy tolerance.
    Returns True if Levenshtein distance <= max_distance.
    """
    query_lower = query_name.lower().strip()
    doc_lower = document_name.lower().strip()
    
    # Exact match
    if query_lower == doc_lower:
        return True
    
    # Check if one contains the other (for partial matches)
    if query_lower in doc_lower or doc_lower in query_lower:
        return True
    
    # Fuzzy match with Levenshtein distance
    distance = levenshtein_distance(query_lower, doc_lower)
    max_allowed = min(max_distance, len(query_lower) // 3)  # Allow up to 1/3 of length
    return distance <= max_allowed


def extract_proper_nouns(text: str) -> List[str]:
    """
    Extract potential proper nouns (capitalized words, likely names).
    Simple heuristic: words that start with capital letters and are not at sentence start.
    """
    import re
    # Find capitalized words (potential names)
    # Pattern: word boundary, capital letter, followed by lowercase letters
    pattern = r'\b[A-Z][a-z]+\b'
    matches = re.findall(pattern, text)
    # Filter out common words that are always capitalized
    common_caps = {'The', 'A', 'An', 'And', 'Or', 'But', 'In', 'On', 'At', 'To', 'For', 'Of', 'With', 'By'}
    proper_nouns = [m for m in matches if m not in common_caps and len(m) > 2]
    return proper_nouns


def detect_name_in_query(query: str) -> Tuple[bool, List[str]]:
    """
    Detect if query contains person/company names.
    Returns (has_name, list_of_names).
    """
    proper_nouns = extract_proper_nouns(query)
    # Heuristic: if we have 2+ capitalized words together, likely a name
    # Or if we have a capitalized word followed by another capitalized word
    import re
    # Pattern for "FirstName LastName" or "Company Name"
    name_pattern = r'\b[A-Z][a-z]+\s+[A-Z][a-z]+\b'
    name_matches = re.findall(name_pattern, query)
    
    all_names = list(set(proper_nouns + name_matches))
    has_name = len(all_names) > 0
    
    return (has_name, all_names)


def classify_query_intent(query: str) -> str:
    """
    Classify query intent: FIND, SUMMARIZE, EXPLAIN, COMPARE, etc.
    """
    q_lower = query.lower()
    
    # FIND intent
    find_patterns = ["find", "search", "locate", "get", "fetch", "retrieve", "show me"]
    if any(pattern in q_lower for pattern in find_patterns):
        return "FIND"
    
    # SUMMARIZE intent
    summarize_patterns = ["summarize", "summarise", "summary", "overview", "brief", "sum up"]
    if any(pattern in q_lower for pattern in summarize_patterns):
        return "SUMMARIZE"
    
    # EXPLAIN intent
    explain_patterns = ["explain", "why", "how does", "how do", "what is", "what are"]
    if any(pattern in q_lower for pattern in explain_patterns):
        return "EXPLAIN"
    
    # COMPARE intent
    compare_patterns = ["compare", "difference", "versus", "vs", "contrast"]
    if any(pattern in q_lower for pattern in compare_patterns):
        return "COMPARE"
    
    # DEFAULT: FIND (most common intent)
    return "FIND"


def extract_source_reference(question: str) -> int | None:
    """
    Extract source number from question (e.g., "source 1", "source [1]", "document 1").
    Returns the source number (1-indexed) or None if not found.
    """
    import re
    q_lower = question.lower()
    # Patterns: "source 1", "source [1]", "document 1", "doc 1", "[1]", etc.
    patterns = [
        r"source\s*\[?(\d+)\]?",
        r"document\s*\[?(\d+)\]?",
        r"doc\s*\[?(\d+)\]?",
        r"\[(\d+)\]",
    ]
    for pattern in patterns:
        match = re.search(pattern, q_lower)
        if match:
            try:
                return int(match.group(1))
            except (ValueError, IndexError):
                continue
    return None


async def extract_search_keywords(user_query: str) -> str:
    """
    Extract core search terms from user query, removing instruction words.
    PRESERVES proper nouns (names) to ensure they're not removed.
    This prevents the vector database from matching on "summarize" instead of "Lily".
    
    Example:
    Input: "Summarize the personal statement for Lily regarding cross-border business"
    Output: "Lily personal statement cross-border business"
    """
    if not user_query:
        return user_query
    
    # Extract proper nouns BEFORE cleaning (to preserve them)
    proper_nouns = extract_proper_nouns(user_query)
    proper_nouns_lower = [pn.lower() for pn in proper_nouns]
    
    if not ANTHROPIC_API_KEY:
        # Fallback: simple keyword extraction using regex
        import re
        # Remove common instruction words
        instruction_patterns = [
            r"\bsummarize\b",
            r"\bsummarise\b",
            r"\btell me about\b",   
            r"\btell me\b",
            r"\bfind\b",
            r"\bsearch for\b",
            r"\bwhat is\b",
            r"\bwhat are\b",
            r"\bwhat does\b",
            r"\bexplain\b",
            r"\bdescribe\b",
            r"\bshow me\b",
            r"\bget\b",
            r"\bfetch\b",
            r"\bretrieve\b",
        ]
        cleaned = user_query
        for pattern in instruction_patterns:
            cleaned = re.sub(pattern, "", cleaned, flags=re.IGNORECASE)
        # Clean up extra spaces
        cleaned = re.sub(r"\s+", " ", cleaned).strip()
        
        # Ensure proper nouns are preserved
        if proper_nouns:
            # Add back any proper nouns that might have been removed
            cleaned_lower = cleaned.lower()
            for pn in proper_nouns:
                if pn.lower() not in cleaned_lower:
                    cleaned = f"{pn} {cleaned}".strip()
        
        return cleaned if cleaned else user_query
    
    try:
        headers = {
            "Content-Type": "application/json",
            "x-api-key": ANTHROPIC_API_KEY,
            "anthropic-version": "2023-06-01",
        }
        url = get_anthropic_api_url()
        
        prompt = f"""Extract the core search terms from this user request. 
Remove instructions like "summarize", "tell me about", "find", "explain", "describe".
Focus on Names, Specific Documents, Topics, and Entities.

User Request: "{user_query}"

Output ONLY the search terms. Do not include explanations or additional text."""

        payload = {
            "model": "claude-3-5-haiku-20241022",  # Use Haiku for cheap, fast extraction
            "max_tokens": 1000,
            "messages": [
                {
                    "role": "user",
                    "content": prompt
                }
            ]
        }
        
        async with httpx.AsyncClient(timeout=5.0) as client:
            res = await client.post(url, headers=headers, json=payload)
            res.raise_for_status()
            data = res.json()
            content = data.get("content", [])
            if isinstance(content, list) and content:
                extracted = content[0].get("text", "").strip()
                if extracted:
                    return extracted
    except Exception as e:
        # Fallback to simple regex extraction on error
        print(f"Keyword extraction failed: {e}")
        import re
        instruction_patterns = [
            r"\bsummarize\b",
            r"\bsummarise\b",
            r"\btell me about\b",
            r"\btell me\b",
            r"\bfind\b",
            r"\bsearch for\b",
            r"\bwhat is\b",
            r"\bwhat are\b",
            r"\bwhat does\b",
            r"\bexplain\b",
            r"\bdescribe\b",
            r"\bshow me\b",
        ]
        cleaned = user_query
        for pattern in instruction_patterns:
            cleaned = re.sub(pattern, "", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\s+", " ", cleaned).strip()
        return cleaned if cleaned else user_query
    
    return user_query

def is_meta_question(question: str) -> bool:
    """
    Detect if question is about capabilities/system (meta) vs document content.
    Meta questions should be answered with general knowledge.
    """
    q_lower = question.lower().strip()
    meta_patterns = [
        "what can you do",
        "what could you do",
        "what are you",
        "what do you do",
        "how do you work",
        "what is your purpose",
        "what are your capabilities",
        "what can you help",
        "how can you help",
        "what features",
        "what functionality",
        "what is orbit ai",
        "who are you",
        "introduce yourself",
        "what is this",
        "what is this system",
        "what is this platform",
    ]
    return any(pattern in q_lower for pattern in meta_patterns)


def has_question_overlap(
    question: str,
    sources: List[AskSource],
    previous_messages: List[ChatMessage] | None = None,
    decisions: List[AskDecision] | None = None,
) -> bool:
    """
    Return True if any meaningful keyword from the question appears in the sources or decisions.
    This is a lightweight guard to avoid hallucinations when sources are unrelated.
    
    IMPORTANT: For follow-up questions, we're more lenient because the user is continuing
    a conversation about a topic that was already validated.
    """
    q_lower = (question or "").lower()
    q_tokens = [t for t in re.split(r"\W+", q_lower) if len(t) > 3]
    
    # CRITICAL: If this is a follow-up question (has previous messages), be MORE lenient
    # The user is continuing a conversation, so we should allow it even if the specific
    # words don't match the sources
    is_followup = previous_messages and len(previous_messages) > 0
    
    # Check for follow-up patterns - these should ALWAYS be allowed if there's history
    followup_patterns = ["yes", "more", "tell", "give", "what", "how", "why", "explain", 
                        "elaborate", "detail", "about", "background", "education", "experience"]
    if is_followup and any(pattern in q_lower for pattern in followup_patterns):
        print(f"[DEBUG] has_question_overlap: ALLOWING follow-up question with history")
        return True
    
    # Extract names from the question - if the question contains a proper name that's in sources, allow it
    name_pattern = r'\b[A-Z][a-z]+\s+[A-Z][a-z]+\b'
    question_names = re.findall(name_pattern, question)
    if question_names and sources:
        source_text = " ".join([f"{s.title or ''} {s.file_name or ''} {s.snippet or ''}".strip() for s in sources]).lower()
        for name in question_names:
            if name.lower() in source_text:
                print(f"[DEBUG] has_question_overlap: ALLOWING because name '{name}' found in sources")
                return True
    
    if previous_messages:
        for msg in previous_messages[-5:]:
            q_tokens.extend(
                [t for t in re.split(r"\W+", (msg.content or "").lower()) if len(t) > 3]
            )
    q_tokens = list(dict.fromkeys(q_tokens))
    
    # Check if question is about decisions
    decision_keywords = ["decision", "decisions", "outcome", "outcomes", "action", "actions", 
                         "invest", "investment", "invested", "passed", "declined", "approved",
                         "rejected", "recent decisions", "decision history", "what decisions"]
    if any(keyword in q_lower for keyword in decision_keywords):
        # If decisions are provided, allow the question
        if decisions and len(decisions) > 0:
            return True
    
    # Check sources
    if sources:
        source_text = " ".join(
            [
                f"{s.title or ''} {s.file_name or ''} {s.snippet or ''}".strip()
                for s in sources
            ]
        ).lower()
        if q_tokens and any(token in source_text for token in q_tokens):
            return True
    
    # Check decisions content if provided
    if decisions and q_tokens:
        decision_text = " ".join(
            [
                f"{d.startup_name or ''} {d.action_type or ''} {d.outcome or ''} {d.notes or ''}".strip()
                for d in decisions
            ]
        ).lower()
        if any(token in decision_text for token in q_tokens):
            return True
    
    print(f"[DEBUG] has_question_overlap: REJECTING - no overlap found. q_tokens={q_tokens[:10]}")
    return False


# System prompt for query contextualization (ChatGPT-style)
CONTEXTUALIZE_SYSTEM_PROMPT = """You are an expert search query generator. 
Your task is to rewrite the "Follow-up Question" into a standalone, specific search query based on the "Chat History".

CRITICAL INSTRUCTIONS:
1. **RESOLVE PRONOUNS**: If the Follow-up Question contains "him", "her", "it", "they", replace it with the specific NAME or ENTITY from the *immediately preceding* User/Assistant exchange.
2. **IGNORE OLD TOPICS**: If the chat started about "Giga Energy" but the last message was about "George", and the user asks "tell me about him", you MUST ask about "George". Ignore Giga Energy.
3. **BE EXPLICIT**: The output must be a full sentence that can be searched in a database.
4. **DO NOT ANSWER**: Do not answer the question. Only output the rewritten query.

---
EXAMPLES (Follow these patterns):

History:
User: Who is Elon Musk?
Assistant: He is the CEO of Tesla.
Follow-up Question: How old is he?
Rewritten Query: How old is Elon Musk?

History:
User: Tell me about the Q3 Report.
Assistant: Here is the Q3 summary...
User: actually, forget that. Who is Sarah Jones?
Assistant: Sarah Jones is the new VP of Sales.
Follow-up Question: Tell me more about her background.
Rewritten Query: Tell me more about Sarah Jones background and resume.

History:
User: Search for Giga Energy.
Assistant: Found 3 documents about Giga Energy.
User: Okay, now look for George Goloborodkin.
Assistant: I found George's resume.
Follow-up Question: What is his email?
Rewritten Query: What is George Goloborodkin's email?
---

Now, rewrite the following:"""


async def rewrite_query_with_llm(question: str, previous_messages: List[ChatMessage] | None = None) -> str:
    """
    ChatGPT-style query rewriting: Use Claude Haiku to contextualize vague follow-up questions.
    
    This is the "Invisible Step" that ChatGPT performs before searching:
    - User: "What do you know about George Goloborodkin?"
    - AI: "George is an intern..."
    - User: "Tell me more about him."
    - System rewrites to: "Tell me more about George Goloborodkin."
    - THEN searches for "George Goloborodkin" (not "him")
    
    This prevents the database from searching for vague words like "him" and returning irrelevant results.
    """
    if not question:
        return question
    
    # If no previous messages, still check if question needs rewriting (might have pronouns from context)
    if not previous_messages:
        previous_messages = []
    
    # Check if question contains pronouns using regex word boundaries (more robust)
    import re
    q_lower = question.lower()
    pronoun_pattern = r'\b(it|its|him|his|her|she|they|them|their|this|that|these|those)\b'
    has_pronouns = bool(re.search(pronoun_pattern, question, re.IGNORECASE))
    
    # Check for affirmative-only responses
    affirmative_only = q_lower.strip() in {
        "yes",
        "yes please",
        "please",
        "ok",
        "okay",
        "sure",
        "go ahead",
        "yep",
        "yeah",
    }
    
    # Check for vague follow-up patterns (case-insensitive, more comprehensive)
    vague_patterns = [
        "tell me more",
        "tell me all",
        "what about",
        "and what",
        "how about",
        "what else",
        "tell more",
        "more about",
        "more details",
        "more info",
        "more information",
        "expand on",
        "elaborate on",
        "go on",
        "all you know",
        "everything about",
        "what's inside",
        "what is inside",
    ]
    has_vague_pattern = any(pattern in q_lower for pattern in vague_patterns)
    
    # ALWAYS rewrite if there's chat history and the question is short/vague
    # This ensures follow-up questions get properly contextualized
    is_short_question = len(question.split()) <= 15  # Increased threshold
    has_chat_history = previous_messages and len(previous_messages) > 0
    
    # Rewrite if: has pronouns, is vague, is affirmative, OR (is short AND has history)
    # ALSO rewrite if question contains "all you know", "everything", "comprehensive" - these need context
    has_comprehensive_intent = any(phrase in q_lower for phrase in ["all you know", "everything", "comprehensive", "all about", "what's inside", "what is inside"])
    
    should_rewrite = has_pronouns or affirmative_only or has_vague_pattern or has_comprehensive_intent or (is_short_question and has_chat_history)
    
    # If no chat history but question has pronouns or is vague, still try to improve it
    if not should_rewrite and not has_chat_history:
        return question
    
    # If we should rewrite but no history, return as-is (can't contextualize without history)
    if should_rewrite and not has_chat_history:
        return question
    
    # CRITICAL: Extract names from ALL messages (user AND assistant) for robust pronoun resolution
    # The full name might appear in the assistant's response, not the user's question
    all_text = " ".join([m.content for m in previous_messages[-6:]])  # Last 6 messages
    all_names = re.findall(r'\b[A-Z][a-z]+\s+[A-Z][a-z]+\b', all_text)  # "FirstName LastName"
    single_names = re.findall(r'\b[A-Z][a-z]{2,}\b', all_text)  # Single capitalized words (potential names)
    # Filter out common words
    common_words = {'The', 'This', 'That', 'Here', 'There', 'What', 'When', 'Where', 'Which', 'Could', 'Would', 'Should', 'Based', 'Found', 'Sorry', 'Please'}
    single_names = [n for n in single_names if n not in common_words]
    
    print(f"[DEBUG] Names found in history - Full names: {all_names}, Single names: {single_names[:5]}")
    
    # Format the history into a clean dialogue string (last 4 messages for better context)
    history_text = ""
    recent_messages = previous_messages[-4:] if len(previous_messages) >= 4 else previous_messages
    for msg in recent_messages:
        role = "User" if msg.role == "user" else "Assistant"
        # Truncate very long messages to avoid token limits
        content = msg.content[:500] if len(msg.content) > 500 else msg.content
        history_text += f"{role}: {content}\n"
    
    # Construct the final input in the format the system prompt expects
    final_user_content = f"History:\n{history_text}Follow-up Question: {question}\nRewritten Query:"
    
    try:
        if not ANTHROPIC_API_KEY:
            # Fallback to simple replacement if no API key
            # Use names extracted from ALL messages (above)
            if (has_pronouns or affirmative_only) and (all_names or single_names):
                # Prefer full names, fall back to single names
                main_subject = all_names[-1] if all_names else single_names[-1]
                print(f"[DEBUG] No API key - using fallback with subject: {main_subject}")
                if has_pronouns:
                    rewritten = question
                    for pronoun in ["him", "her", "it", "they", "them", "his", "her", "their", "this", "that"]:
                        rewritten = re.sub(rf'\b{pronoun}\b', main_subject, rewritten, flags=re.IGNORECASE)
                    return rewritten
                if affirmative_only:
                    return f"Tell me more about {main_subject}"
            return question
        
        headers = {
            "Content-Type": "application/json",
            "x-api-key": ANTHROPIC_API_KEY,
            "anthropic-version": "2023-06-01",
        }
        url = get_anthropic_api_url()
        
        # Use system message + user message (ChatGPT-style)
        payload = {
            "model": "claude-3-5-haiku-20241022",  # Use Haiku for cheap, fast rewriting
            "max_tokens": 100,  # Reduced since we only want the query
            "system": CONTEXTUALIZE_SYSTEM_PROMPT,
            "messages": [
                {
                    "role": "user",
                    "content": final_user_content
                }
            ]
        }
        
        async with httpx.AsyncClient(timeout=10.0) as client:
            res = await client.post(url, headers=headers, json=payload)
            res.raise_for_status()
            data = res.json()
            content = data.get("content", [])
            if isinstance(content, list) and content:
                rewritten = content[0].get("text", "").strip()
                # Clean up any prefixes or extra text the model might add
                rewritten = rewritten.lstrip("Rewritten Query:").lstrip("Query:").lstrip("-").strip()
                # Remove any explanatory text after the question (look for newlines)
                if "\n" in rewritten:
                    rewritten = rewritten.split("\n")[0].strip()
                # Remove trailing periods if it's just a period
                rewritten = rewritten.rstrip(".")
                
                if rewritten and rewritten != question:
                    print(f"[DEBUG] Query rewritten by LLM: '{question}' -> '{rewritten}'")
                    # Validate: if original had "him" and rewritten doesn't contain any name from history, force fix it
                    if has_pronouns and (all_names or single_names):
                        # Check if rewritten contains any of the names we found in history
                        rewritten_lower = rewritten.lower()
                        has_name_from_history = any(name.lower() in rewritten_lower for name in all_names) or \
                                               any(name.lower() in rewritten_lower for name in single_names[:5])
                        
                        if not has_name_from_history:
                            print(f"[DEBUG] ⚠️ WARNING: LLM rewrite doesn't contain names from history, using fallback")
                            # Use the most recent full name, or fall back to single name
                            main_subject = all_names[-1] if all_names else single_names[-1] if single_names else None
                            if main_subject:
                                print(f"[DEBUG] Using fallback: replacing pronouns with '{main_subject}'")
                                rewritten = question
                                for pronoun in ["him", "her", "it", "they", "them", "his", "her", "their", "this", "that"]:
                                    rewritten = re.sub(rf'\b{pronoun}\b', main_subject, rewritten, flags=re.IGNORECASE)
                                return rewritten
                    return rewritten
    except Exception as e:
        # Fallback to simple replacement on error
        print(f"[DEBUG] Query rewriting exception: {e}")
        # Use names extracted from ALL messages (extracted above)
        if has_pronouns and (all_names or single_names):
            main_subject = all_names[-1] if all_names else single_names[-1]
            print(f"[DEBUG] Exception fallback: replacing pronouns with '{main_subject}'")
            rewritten = question
            for pronoun in ["him", "her", "it", "they", "them", "his", "her", "their", "this", "that"]:
                rewritten = re.sub(rf'\b{pronoun}\b', main_subject, rewritten, flags=re.IGNORECASE)
            return rewritten
        elif affirmative_only and (all_names or single_names):
            main_subject = all_names[-1] if all_names else single_names[-1]
            return f"Tell me more about {main_subject}"
    
    return question


def resolve_followup_context(question: str, previous_messages: List[ChatMessage] | None = None) -> str:
    """
    Legacy synchronous wrapper. For async contexts, use rewrite_query_with_llm instead.
    """
    if not question or not previous_messages:
        return question
    q_lower = question.lower()
    affirmative_only = q_lower.strip() in {
        "yes",
        "yes please",
        "please",
        "ok",
        "okay",
        "sure",
        "go ahead",
    }
    if affirmative_only:
        last_user = next((m.content for m in reversed(previous_messages) if m.role == "user"), "").strip()
        if last_user:
            return f"{last_user}\n\nFollow-up request: Provide a more complete answer from the available sources."
        return question
    if " it " not in f" {q_lower} " and " it?" not in q_lower and " it." not in q_lower:
        return question
    last_user = next((m.content for m in reversed(previous_messages) if m.role == "user"), "").strip()
    if not last_user:
        return question
    return f"{last_user}\n\nFollow-up question: {question}"


def build_answer_prompt(question: str, sources: List[AskSource], decisions: List[AskDecision], previous_messages: List[ChatMessage] = None) -> str:
    is_meta = is_meta_question(question)
    is_comprehensive = is_comprehensive_question(question)
    source_ref = extract_source_reference(question)
    safe_sources = (sources or [])[:ASK_MAX_SOURCES]
    max_snippet_chars = ASK_MAX_SNIPPET_CHARS
    if is_comprehensive:
        max_snippet_chars = max(ASK_MAX_SNIPPET_CHARS, 2000)  # Much larger for comprehensive questions
    source_lines: List[str] = []
    for idx, src in enumerate(safe_sources, start=1):
        title = src.title or src.file_name or f"Source {idx}"
        snippet = (src.snippet or "").strip()
        # If user references a specific source, give it much more space
        if source_ref == idx:
            max_snippet_chars_for_this = max(max_snippet_chars, 3000)  # Even larger for referenced source
        else:
            max_snippet_chars_for_this = max_snippet_chars
        if len(snippet) > max_snippet_chars_for_this:
            snippet = snippet[:max_snippet_chars_for_this] + "…"
        source_lines.append(f"[{idx}] {title}\n{snippet}")

    decision_lines: List[str] = []
    for d in decisions or []:
        summary = " | ".join(
            [part for part in [d.startup_name, d.action_type, d.outcome, d.notes] if part]
        )
        if summary:
            decision_lines.append(f"- {summary}")

    sources_block = "\n\n".join(source_lines) if source_lines else "No sources available."
    decisions_block = "\n".join(decision_lines) if decision_lines else "No decision history available."
    
    # Build conversation history context - INCLUDE ALL MESSAGES for full context
    # BUT emphasize MOST RECENT messages for pronoun resolution
    conversation_context = ""
    if previous_messages and len(previous_messages) > 0:
        # Include ALL messages (up to 20 to avoid token limits, but prioritize recent)
        recent_messages = previous_messages[-20:] if len(previous_messages) > 20 else previous_messages
        conversation_lines = []
        for i, msg in enumerate(recent_messages):
            role_label = "User" if msg.role == "user" else "Assistant"
            # Truncate very long messages to avoid token limits
            content = msg.content[:1000] if len(msg.content) > 1000 else msg.content
            # Mark most recent messages (last 3) as MOST RECENT
            if i >= len(recent_messages) - 3:
                conversation_lines.append(f"{role_label} (MOST RECENT): {content}")
            else:
                conversation_lines.append(f"{role_label}: {content}")
        
        # Get the last user question explicitly
        last_user_q = next((m.content for m in reversed(recent_messages) if m.role == "user"), "")
        
        conversation_context = f"\n\n=== PREVIOUS CONVERSATION HISTORY (USE THIS TO UNDERSTAND PRONOUNS AND CONTEXT) ===\n" + \
                             f"⚠️ IMPORTANT: Messages marked 'MOST RECENT' are the most recent conversation. If user says 'him', 'her', 'it', check the MOST RECENT messages first!\n" + \
                             f"Last User Question: {last_user_q[:200] if last_user_q else 'N/A'}\n" + \
                             f"\n" + "\n".join(conversation_lines) + "\n=== END OF CONVERSATION HISTORY ===\n"
        # Debug logging - DETAILED
        print(f"[DEBUG] ✅ Conversation history included in prompt: {len(recent_messages)} messages")
        print(f"[DEBUG] First message: {recent_messages[0].role}: {recent_messages[0].content[:150]}")
        print(f"[DEBUG] Last message: {recent_messages[-1].role}: {recent_messages[-1].content[:150]}")
        # Check if history contains names that might be referenced
        all_content = " ".join([m.content for m in recent_messages])
        import re
        names = re.findall(r'\b[A-Z][a-z]+\s+[A-Z][a-z]+\b', all_content)
        if names:
            print(f"[DEBUG] Found names in history: {set(names)}")
    else:
        print(f"[DEBUG] ❌❌❌ CRITICAL WARNING: No conversation history provided! previous_messages={previous_messages}")
        print(f"[DEBUG] This means pronouns like 'him', 'her' cannot be resolved!")
    
    
    is_raw_text = is_raw_text_request(question)
    
    if is_meta:
        # Meta questions: answer with general knowledge about Orbit AI capabilities
        return f"""You are Orbit AI, a VC intelligence system built for investment teams. Answer this question about your capabilities and features.

Question:
{question}

Answer based on what Orbit AI can do:
- Answer questions about uploaded documents (pitch decks, memos, meeting notes)
- Extract structured information from unstructured documents
- Track investment decisions and outcomes
- Provide insights from your fund's knowledge base
- Search across all uploaded sources semantically
- Help with due diligence by finding relevant information quickly

Be helpful and specific. Explain what you can do and how you help investment teams.
"""
    else:
        # Document questions: use sources only
        comprehensive_instruction = ""
        if is_comprehensive:
            comprehensive_instruction = "\n\n🚨 CRITICAL: The user is asking for a COMPREHENSIVE answer. This means:\n- Provide ALL available information from the sources about this topic\n- Be thorough, detailed, and include all relevant details\n- Don't summarize or be brief - give the FULL picture\n- Include all sections, data points, qualifications, experiences, and any other information\n- If asked about a person (e.g., 'all you know about George'), include everything: background, education, experience, role, responsibilities, etc.\n- If asked about a document (e.g., 'what's inside source 1'), provide a complete overview of all content\n- Do NOT say 'limited information' or 'very limited' - extract and present EVERYTHING that exists in the sources\n- Be exhaustive, not defensive"
        raw_text_instruction = ""
        if is_raw_text:
            raw_text_instruction = "\n\nIMPORTANT: The user is asking for RAW/EXACT TEXT. Provide the source snippets verbatim (no paraphrasing). If the text is truncated, say so explicitly. Preserve formatting and line breaks when possible."
        
        # Build source reference instruction if user mentions a specific source
        source_ref_instruction = ""
        if source_ref:
            source_ref_instruction = f"\n\nIMPORTANT: The user is asking about SOURCE [{source_ref}]. Focus primarily on that source and provide a COMPREHENSIVE overview of everything in that document. Include all key details, sections, data points, and information from source [{source_ref}]."
        
        # Build the prompt with conversation history at the top
        history_section = conversation_context if conversation_context else "\n\n=== PREVIOUS CONVERSATION HISTORY ===\n(No previous conversation history available)\n=== END OF CONVERSATION HISTORY ===\n"
        
        return f"""You are Orbit AI, a VC intelligence system. You answer questions based on the provided sources and conversation history.

{history_section}

🚨 CRITICAL: READ THE CONVERSATION HISTORY ABOVE CAREFULLY BEFORE ANSWERING!

CRITICAL RULES:
1. **MANDATORY: CHECK CONVERSATION HISTORY FIRST**. If the user uses pronouns like "him", "her", "it", "they", "them", "his", "her", "their", "this", "that", "these", "those", you MUST look in the conversation history above to find what they're referring to. The conversation history shows the full context of what was discussed previously.
2. **IF YOU SEE "tell me more about him" AND THE HISTORY SHOWS A PREVIOUS QUESTION ABOUT "George Goloborodkin"**, then "him" = "George Goloborodkin". Use the conversation history to resolve ALL pronouns. **NEVER say "I cannot determine who 'him' refers to" if there is conversation history above - ALWAYS check it first!**
3. If the user asks "what's inside", "what is in source X", "all you know", or similar questions about document contents, provide a COMPREHENSIVE and DETAILED answer covering ALL information in the relevant source(s). Do NOT be brief or defensive - give the FULL picture.
4. If the user references a specific source (e.g., "source 1", "source [1]", "document 1"), focus on that source and provide comprehensive details from it. Recognize that [1] refers to the first source, [2] to the second, etc.
5. The sources provided may NOT be relevant to the question. You MUST verify relevance before answering.
6. If the sources DO contain relevant details that DIRECTLY answer the question, provide a thorough, well-structured answer using those details. Be comprehensive and include all relevant information from the sources.
7. If the sources do NOT contain relevant information about the question topic, you MUST say: "I don't have information about this in the provided sources. Please upload relevant documents or try a different question."
8. Do NOT answer with information that is tangentially related but doesn't actually address the question.
9. If a source talks about a completely different topic (e.g., trading/ATR when asked about a person's resume), you MUST reject it and say you don't have information.
10. Cite sources using [1], [2], etc. for every claim.
11. Do NOT be overly apologetic. If you have information, present it confidently and thoroughly. Only apologize if you truly have no relevant information.
12. When the user asks about something mentioned in the conversation (e.g., "tell me more about him" after discussing George), search the sources for information about that person/entity, even if the current question is vague.{comprehensive_instruction}{raw_text_instruction}{source_ref_instruction}

Answer style:
- Prioritize comprehensive, coherent narrative answers grounded in sources.
- Prefer completeness over brevity when sources list multiple items.
- Do not force a sectioned structure; use paragraphs with bullets only when they improve clarity.
- If information is sparse, still provide the most complete answer possible from the available evidence.
- Use bullet points for responsibilities, qualifications, and scope when asked.
- For comprehensive questions, expand with all relevant details from sources in a flowing, human-like summary.
- For raw text requests, return verbatim snippets with source labels and no paraphrasing.

Question:
{question}

Sources:
{sources_block}

Decision history (optional context):
{decisions_block}

Remember: 
- ALWAYS check the conversation history above to understand pronouns and context
- If the answer isn't in the sources above, you MUST say you don't have that information. Never make up answers.
"""

# Fast model for simple questions (3-5x faster)
HAIKU_MODEL = "claude-3-5-haiku-20241022"

def is_simple_question(question: str, sources: List[AskSource]) -> bool:
    """
    Detect if question is simple enough for Haiku (3-5x faster).
    Simple = short question, few sources, straightforward answer expected.
    """
    # Simple heuristics:
    # 1. Short question (< 15 words)
    word_count = len(question.split())
    if word_count > 15:
        return False
    
    # 2. Few sources (1-2)
    if len(sources) > 2:
        return False
    
    # 3. Not asking for analysis/comparison
    complex_keywords = ["compare", "analyze", "why", "explain", "evaluate", "assess", "strategy"]
    question_lower = question.lower()
    if any(kw in question_lower for kw in complex_keywords):
        return False
    
    # 4. Asking for facts (what, who, when, where, how much)
    simple_patterns = ["what is", "what are", "who is", "when", "where", "how much", "how many", "tell me about"]
    if any(pattern in question_lower for pattern in simple_patterns):
        return True
    
    # Default: use Sonnet for better quality
    return False


async def call_anthropic_answer(prompt: str, question: str = "", sources: List[AskSource] = None) -> str:
    if not ANTHROPIC_API_KEY:
        raise HTTPException(
            status_code=503,
            detail="ANTHROPIC_API_KEY not set. Set it in the server environment to use Claude."
        )

    headers = {
        "Content-Type": "application/json",
        "x-api-key": ANTHROPIC_API_KEY,
        "anthropic-version": "2023-06-01",
    }
    url = get_anthropic_api_url()
    default_url = "https://api.anthropic.com/v1/messages"

    # Choose model based on question complexity (Haiku is 3-5x faster)
    use_haiku = question and sources and is_simple_question(question, sources)
    model_list = ([HAIKU_MODEL] + ANTHROPIC_MODEL_FALLBACKS) if use_haiku else ANTHROPIC_MODEL_FALLBACKS
    # Set max_tokens appropriately (tokens are cheap, user trust is expensive)
    max_tokens = 10000 if use_haiku else ASK_MAX_TOKENS
    
    last_error: Optional[str] = None
    async with httpx.AsyncClient(timeout=60.0) as client:
        for model_name in [m for m in model_list if m]:
            payload = {
                "model": model_name,
                "max_tokens": max_tokens,
                "temperature": 0.5,
                "system": "You are Orbit AI, a VC intelligence system. You answer questions STRICTLY from provided sources only. Never use general knowledge. If information isn't in the sources, say so explicitly. Always cite sources with [1], [2], etc.",
                "messages": [
                    {"role": "user", "content": prompt}
                ],
            }

            res = await client.post(url, headers=headers, json=payload)
            if res.status_code == 404 and url != default_url:
                res = await client.post(default_url, headers=headers, json=payload)
            if res.status_code == 404 and is_model_not_found(res):
                last_error = f"Model not found: {model_name}"
                continue
            if res.status_code >= 400:
                body = res.text[:400].strip()
                raise HTTPException(
                    status_code=502,
                    detail=f"Claude API error ({res.status_code}): {body or 'empty response'}",
                )

            data = res.json() or {}
            content = data.get("content") or []
            text = ""
            if isinstance(content, list) and content:
                text = content[0].get("text") or ""
            elif isinstance(content, str):
                text = content
            if not text:
                raise HTTPException(status_code=502, detail="Claude returned empty content.")
            return text.strip()

    raise HTTPException(status_code=503, detail=last_error or "No Claude model available.")

def normalize_startup_data(data: Dict[str, Any]) -> StartupData:
    """Normalize extracted startup data to match schema"""
    def safe_str(val: Any) -> str:
        return val.strip() if isinstance(val, str) else (str(val).strip() if val is not None else "")
    def safe_int(val: Any, default: int = 0) -> int:
        try:
            if val is None:
                return default
            if isinstance(val, (int, float)):
                return int(val)
            if isinstance(val, str):
                cleaned = re.sub(r'[^\d.]', '', val)
                return int(float(cleaned)) if cleaned else default
            return default
        except Exception:
            return default

    # Handle geoMarkets (accept snake_case + common synonyms like region)
    geo_markets = data.get('geoMarkets', data.get('geo_markets', data.get('region', data.get('regions', data.get('geography', [])))))
    if isinstance(geo_markets, str):
        geo_markets = [g.strip() for g in re.split(r'[,;|]', geo_markets)]
    
    # Handle fundingTarget
    funding_target = data.get('fundingTarget', data.get('funding_target', 0))
    if isinstance(funding_target, str):
        # Extract number from string
        funding_target = re.sub(r'[^\d.]', '', funding_target)
        funding_target = int(float(funding_target)) if funding_target else 0
    funding_target = safe_int(funding_target, 0)
    
    return StartupData(
        companyName=safe_str(data.get('companyName', data.get('company_name', data.get('name', '')))),
        geoMarkets=geo_markets if isinstance(geo_markets, list) else [],
        industry=safe_str(data.get('industry', data.get('sector', data.get('startup_industry', '')))),
        fundingTarget=safe_int(funding_target, 0),
        fundingStage=safe_str(data.get('fundingStage', data.get('funding_stage', data.get('stage', '')))),
        availabilityStatus='present'
    )

def normalize_investor_data(data: Dict[str, Any]) -> InvestorData:
    """Normalize extracted investor data to match schema"""
    def safe_str(val: Any) -> str:
        return val.strip() if isinstance(val, str) else (str(val).strip() if val is not None else "")
    def safe_int(val: Any, default: int = 0) -> int:
        try:
            if val is None:
                return default
            if isinstance(val, (int, float)):
                return int(val)
            if isinstance(val, str):
                cleaned = re.sub(r'[^\d.]', '', val)
                return int(float(cleaned)) if cleaned else default
            return default
        except Exception:
            return default

    # Handle lists
    def parse_list(value):
        if isinstance(value, list):
            return value
        if isinstance(value, str):
            return [item.strip() for item in re.split(r'[,;|]', value) if item.strip()]
        return []
    
    # Handle numbers
    def parse_number(value):
        if value is None:
            return 0
        if isinstance(value, (int, float)):
            return int(value)
        if isinstance(value, str):
            # Handle currency and multipliers
            val = value.upper()
            multiplier = 1
            if 'M' in val or 'MILLION' in val:
                multiplier = 1000000
            elif 'K' in val or 'THOUSAND' in val:
                multiplier = 1000
            digits = re.sub(r'[^\d.]', '', val)
            try:
                return int(float(digits) * multiplier) if digits else 0
            except Exception:
                return 0
        # Fallback for any other type
        return safe_int(value, 0)
    
    geo_focus = parse_list(
        data.get('geoFocus') or
        data.get('geo_focus') or 
        data.get('Main Geographies Targeted (labels)') or  # Orbit CSV format
        data.get('Location') or
        data.get('geoMarkets') or 
        data.get('region') or 
        data.get('regions') or 
        data.get('geography') or 
        []
    )
    
    industry_prefs = parse_list(
        data.get('industryPreferences') or 
        data.get('industry_preferences') or 
        data.get('[BD] Vertical Interests / Vertical (labels)') or  # Orbit CSV format
        data.get('Vertical Interests') or
        data.get('industries') or 
        []
    )
    
    stage_prefs = parse_list(
        data.get('stagePreferences') or 
        data.get('stage_preferences') or 
        data.get('stages') or 
        []
    )
    
    # Handle cheque/ticket size - may be a range like "100K-500K" or ">1M"
    cheque_size_raw = data.get('[INV] Cheque Size (labels)') or data.get('Cheque Size') or data.get('Check Size') or ''
    
    if cheque_size_raw and isinstance(cheque_size_raw, str):
        # Parse ranges like "100K-500K" or ">1M"
        if '-' in cheque_size_raw:
            parts = cheque_size_raw.split('-')
            min_ticket = parse_number(parts[0]) if len(parts) > 0 else 0
            max_ticket = parse_number(parts[1]) if len(parts) > 1 else min_ticket * 10
        elif '>' in cheque_size_raw:
            min_ticket = parse_number(cheque_size_raw.replace('>', ''))
            max_ticket = min_ticket * 10
        elif '<' in cheque_size_raw:
            max_ticket = parse_number(cheque_size_raw.replace('<', ''))
            min_ticket = max_ticket // 10
        else:
            min_ticket = parse_number(cheque_size_raw)
            max_ticket = min_ticket * 5
    else:
        min_ticket = parse_number(data.get('minTicketSize') or data.get('min_ticket_size') or data.get('minInvestment') or 0)
        max_ticket = parse_number(data.get('maxTicketSize') or data.get('max_ticket_size') or data.get('maxInvestment') or 10000000)
    
    total_slots = safe_int(data.get('totalSlots') or data.get('total_slots') or data.get('slots') or 3, 3)

    # Handle various column name formats from different sources
    firm_name = safe_str(
        data.get('firmName') or 
        data.get('firm_name') or 
        data.get('Investor name') or  # Orbit CSV format
        data.get('name') or 
        data.get('firm') or 
        data.get('Company Name') or
        ''
    )
    
    member_name = safe_str(
        data.get('memberName') or 
        data.get('member_name') or 
        data.get('🦅 [INV] Team Member (users)') or  # Orbit CSV format
        data.get('[INV] Team Member (users)') or
        data.get('Team Member') or
        data.get('investment_member') or 
        data.get('investorMemberName') or 
        data.get('contactName') or 
        data.get('partnerName') or 
        data.get('personName') or 
        ''
    )
    
    return InvestorData(
        firmName=firm_name,
        memberName=member_name,
        geoFocus=geo_focus,
        industryPreferences=industry_prefs,
        stagePreferences=stage_prefs,
        minTicketSize=min_ticket,
        maxTicketSize=max_ticket,
        totalSlots=total_slots,
        tableNumber=data.get('tableNumber', data.get('table_number', data.get('table', None))),
        availabilityStatus='present'
    )

def normalize_mentor_data(data: Dict[str, Any]) -> MentorData:
    """Normalize extracted mentor data to match schema"""
    def safe_str(val: Any) -> str:
        return val.strip() if isinstance(val, str) else (str(val).strip() if val is not None else "")
    
    def parse_list(value):
        if isinstance(value, list):
            return value
        if isinstance(value, str):
            return [item.strip() for item in re.split(r'[,;|]', value) if item.strip()]
        return []
    
    full_name = safe_str(
        data.get('fullName') or 
        data.get('full_name') or 
        data.get('Full Name') or 
        data.get('name') or 
        ''
    )
    
    email = safe_str(
        data.get('email') or 
        data.get('Email') or 
        ''
    )
    
    linkedin_url = safe_str(
        data.get('linkedinUrl') or 
        data.get('linkedin_url') or 
        data.get('LinkedIn URL') or 
        data.get('LinkedIn') or 
        None
    )
    
    geo_focus = parse_list(
        data.get('geoFocus') or
        data.get('geo_focus') or 
        data.get('Location') or
        data.get('region') or 
        []
    )
    
    industry_prefs = parse_list(
        data.get('industryPreferences') or 
        data.get('industry_preferences') or 
        data.get('Industry Preferences') or
        data.get('industries') or 
        []
    )
    
    expertise_areas = parse_list(
        data.get('expertiseAreas') or 
        data.get('expertise_areas') or 
        data.get('Expertise Areas') or
        data.get('expertise') or 
        []
    )
    
    total_slots = int(data.get('totalSlots') or data.get('total_slots') or data.get('Total Slots') or 3)
    
    return MentorData(
        fullName=full_name,
        email=email,
        linkedinUrl=linkedin_url,
        geoFocus=geo_focus,
        industryPreferences=industry_prefs,
        expertiseAreas=expertise_areas,
        totalSlots=total_slots,
        availabilityStatus='present'
    )

def normalize_corporate_data(data: Dict[str, Any]) -> CorporateData:
    """Normalize extracted corporate data to match schema"""
    def safe_str(val: Any) -> str:
        return val.strip() if isinstance(val, str) else (str(val).strip() if val is not None else "")
    
    def parse_list(value):
        if isinstance(value, list):
            return value
        if isinstance(value, str):
            return [item.strip() for item in re.split(r'[,;|]', value) if item.strip()]
        return []
    
    firm_name = safe_str(
        data.get('firmName') or 
        data.get('firm_name') or 
        data.get('Company Name') or 
        data.get('companyName') or 
        data.get('name') or 
        ''
    )
    
    contact_name = safe_str(
        data.get('contactName') or 
        data.get('contact_name') or 
        data.get('Contact Name') or 
        ''
    )
    
    email = safe_str(
        data.get('email') or 
        data.get('Email') or 
        None
    )
    
    geo_focus = parse_list(
        data.get('geoFocus') or
        data.get('geo_focus') or 
        data.get('Location') or
        data.get('region') or 
        []
    )
    
    industry_prefs = parse_list(
        data.get('industryPreferences') or 
        data.get('industry_preferences') or 
        data.get('Industry Preferences') or
        data.get('industries') or 
        []
    )
    
    partnership_types = parse_list(
        data.get('partnershipTypes') or 
        data.get('partnership_types') or 
        data.get('Partnership Types') or
        []
    )
    
    stages = parse_list(
        data.get('stages') or 
        data.get('Stages') or 
        []
    )
    
    total_slots = int(data.get('totalSlots') or data.get('total_slots') or data.get('Total Slots') or 3)
    
    return CorporateData(
        firmName=firm_name,
        contactName=contact_name,
        email=email,
        geoFocus=geo_focus,
        industryPreferences=industry_prefs,
        partnershipTypes=partnership_types,
        stages=stages,
        totalSlots=total_slots,
        availabilityStatus='present'
    )

async def extract_text_content(file: UploadFile) -> Tuple[str, str]:
    """
    Shared helper to read an uploaded file and extract text_content with best-effort parsing.
    Returns (file_ext, text_content).
    """
    # Read the uploaded file bytes.
    # On some setups, UploadFile may be at EOF (e.g. if something already read the stream),
    # so we retry once after seeking back to the start.
    content = await file.read()
    if not content or len(content) == 0:
        try:
            await file.seek(0)
            content = await file.read()
        except Exception:
            # If seek/read fails, we fall through to the empty-upload guard below.
            pass
    file_ext = file.filename.split('.')[-1].lower() if file.filename else ""
    text_content = None  # Initialize to None

    # Guard: empty upload (common when the browser upload failed or the file is zero bytes)
    if not content or len(content) == 0:
        raise HTTPException(
            status_code=400,
            detail=(
                "Uploaded file is empty (0 bytes). Re-upload the file. "
                "If this keeps happening, the browser may be sending an empty payload or the file may be corrupt. "
                f"filename={file.filename!r}, content_type={getattr(file, 'content_type', None)!r}"
            )
        )

    # Normalize extension and detect by magic bytes only if we don't already recognize a handled type.
    handled_exts = {'pdf', 'xlsx', 'xls', 'csv', 'txt', 'json', 'doc', 'docx'}
    # Fallback detection by magic bytes (override when extension is missing or wrong)
    # NOTE: some PDFs may have leading bytes before the %PDF header, so we search the first chunk.
    head = content[:2048]
    if b'%PDF' in head and file_ext not in handled_exts:
        file_ext = 'pdf'
    elif head[:2] == b'PK' and b'[Content_Types].xml' in head:
        # Peek inside the zip to disambiguate docx vs xlsx
        try:
            import zipfile
            from io import BytesIO
            with zipfile.ZipFile(BytesIO(content)) as z:
                names = set(z.namelist())
                if 'word/document.xml' in names:
                    file_ext = 'docx'
                elif 'xl/workbook.xml' in names or any(n.startswith('xl/') for n in names):
                    file_ext = 'xlsx'
                elif file_ext not in handled_exts:
                    # default fallback
                    file_ext = 'xlsx'
        except Exception:
            # fall back to extension if zip probe fails
            if file_ext not in handled_exts and file_ext not in ['docx', 'xlsx']:
                file_ext = 'xlsx'
    elif content.startswith(b'\xd0\xcf\x11\xe0') and file_ext not in handled_exts:
        # Old Office formats (could be doc or xls); if extension says doc, keep doc, else assume xls
        file_ext = 'doc' if file_ext == 'doc' else 'xls'
    
    # Handle Excel files (XLSX, XLS)
    if file_ext in ['xlsx', 'xls']:
        try:
            if file_ext == 'xlsx':
                try:
                    import openpyxl
                    from io import BytesIO
                    
                    excel_file = BytesIO(content)
                    workbook = openpyxl.load_workbook(excel_file, data_only=True)
                    text_content = ""
                    
                    # Get the first sheet
                    sheet = workbook.active
                    
                    # Extract headers
                    if sheet.max_row > 0:
                        headers = []
                        for cell in sheet[1]:
                            headers.append(str(cell.value) if cell.value else "")
                        text_content += ",".join(headers) + "\n"
                        
                        # Extract data rows
                        for row in sheet.iter_rows(min_row=2, values_only=False):
                            row_data = []
                            for cell in row:
                                row_data.append(str(cell.value) if cell.value else "")
                            text_content += ",".join(row_data) + "\n"
                    
                    if not text_content.strip():
                        raise HTTPException(status_code=400, detail="Excel file appears to be empty.")
                except ImportError:
                    raise HTTPException(
                        status_code=500,
                        detail="XLSX support requires openpyxl. Install with: pip install openpyxl"
                    )
            else:  # XLS
                try:
                    import xlrd
                    
                    workbook = xlrd.open_workbook(file_contents=content)
                    text_content = ""
                    
                    # Get the first sheet
                    sheet = workbook.sheet_by_index(0)
                    
                    # Extract headers
                    if sheet.nrows > 0:
                        headers = [str(sheet.cell_value(0, col)) for col in range(sheet.ncols)]
                        text_content += ",".join(headers) + "\n"
                        
                        # Extract data rows
                        for row_idx in range(1, sheet.nrows):
                            row_data = [str(sheet.cell_value(row_idx, col)) for col in range(sheet.ncols)]
                            text_content += ",".join(row_data) + "\n"
                    
                    if not text_content.strip():
                        raise HTTPException(status_code=400, detail="Excel file appears to be empty.")
                except ImportError:
                    raise HTTPException(
                        status_code=500,
                        detail="XLS support requires xlrd. Install with: pip install xlrd"
                    )
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to parse Excel file: {str(e)}")
        
        # After Excel parsing, text_content should be set
        if text_content is None or not text_content.strip():
            raise HTTPException(status_code=500, detail="Excel file parsing completed but no content extracted.")
    
    # Handle DOCX files (prefer python-docx for tables; fallback to raw XML)
    elif file_ext == 'docx':
        try:
            from io import BytesIO
            try:
                from docx import Document  # type: ignore
                doc = Document(BytesIO(content))
                parts = []
                for p in doc.paragraphs:
                    if p.text and p.text.strip():
                        parts.append(p.text)
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text is not None)
                        if row_text.strip():
                            parts.append(row_text)
                text_content = "\n".join(parts)
            except ImportError:
                # Fallback: manual XML strip
                import zipfile
                with zipfile.ZipFile(BytesIO(content)) as z:
                    with z.open('word/document.xml') as doc_xml:
                        raw_xml = doc_xml.read().decode('utf-8', errors='ignore')
                        # Replace paragraph boundaries with newline
                        raw_xml = raw_xml.replace('</w:p>', '\n')
                        # Strip XML tags
                        text_content = re.sub(r'<[^>]+>', '', raw_xml)
            if not text_content or not text_content.strip():
                raise HTTPException(status_code=400, detail="DOCX appears to have no extractable text. If this is a scanned/image DOCX, re-save as PDF or CSV.")
        except KeyError:
            raise HTTPException(status_code=400, detail="DOCX is missing document.xml. Please re-save the file or export to PDF.")
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Could not extract text from DOCX. Please export to PDF or CSV. Error: {str(e)}")

    # Handle DOC files (legacy binary) — best-effort text extraction; if empty, ask user to convert
    elif file_ext == 'doc':
        try:
            # DOC is legacy OLE; we don't depend on heavy converters here.
            # Best-effort: decode as latin-1 ignoring errors and strip control chars.
            raw = content.decode('latin-1', errors='ignore')
            # Remove nulls and most control chars
            cleaned = re.sub(r'[\x00-\x08\x0B-\x1F\x7F]', ' ', raw)
            # Collapse whitespace
            cleaned = re.sub(r'\s+', ' ', cleaned).strip()
            if len(cleaned) < 20:
                raise HTTPException(status_code=400, detail="DOC (legacy Word) has no extractable text. Please re-save as DOCX or PDF and re-upload.")
            text_content = cleaned
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"DOC (legacy Word) could not be read. Please re-save as DOCX or PDF and re-upload. Error: {str(e)}")

    # Handle PDF files
    elif file_ext == 'pdf':
        # Quick sanity-check: real PDFs should contain a %PDF header near the beginning.
        # Some PDFs may have leading bytes, so search the first chunk instead of only prefix.
        if b"%PDF" not in content[:8192]:
            head_hex = content[:32].hex()
            raise HTTPException(
                status_code=400,
                detail=f'File has ".pdf" extension but does not look like a valid PDF (missing %PDF header). First bytes (hex): {head_hex}'
            )

        # 1) Try PyMuPDF first with PARALLEL processing (often more robust than PyPDF2/pdfplumber on quirky PDFs)
        pymupdf_error = None
        try:
            import fitz  # PyMuPDF
            from io import BytesIO
            import concurrent.futures

            doc = fitz.open(stream=BytesIO(content).getvalue(), filetype="pdf")
            page_limit = min(doc.page_count, MAX_PDF_PAGES)
            
            # Parallel page extraction for speed (using thread pool since PyMuPDF is sync)
            import concurrent.futures
            
            def extract_page_sync(i: int) -> str:
                try:
                    page = doc.load_page(i)
                    page_text = page.get_text("text") or ""
                    return f"\n--- Page {i + 1} ---\n{page_text}"
                except Exception as e:
                    return f"\n--- Page {i + 1} (error: {e}) ---\n"
            
            # Process pages in parallel using ThreadPoolExecutor
            parts = []
            with concurrent.futures.ThreadPoolExecutor(max_workers=MAX_PARALLEL_PAGES) as executor:
                futures = [executor.submit(extract_page_sync, i) for i in range(page_limit)]
                for future in concurrent.futures.as_completed(futures):
                    parts.append(future.result())
            
            # Sort parts by page number to maintain order
            parts.sort(key=lambda x: int(re.search(r'Page (\d+)', x).group(1)) if re.search(r'Page (\d+)', x) else 0)
            
            text_content = "\n".join(parts).strip()
        except Exception as e:
            pymupdf_error = e
            text_content = None

        # If PyMuPDF extracted real text, we're done.
        if text_content and len(text_content.strip()) >= 50:
            return file_ext, text_content
        
        # 2) If PyMuPDF extracted <50 chars, likely scanned/image PDF - try Claude Vision
        if text_content and len(text_content.strip()) < 50 and ANTHROPIC_API_KEY:
            try:
                import fitz  # PyMuPDF
                from io import BytesIO
                
                doc = fitz.open(stream=BytesIO(content).getvalue(), filetype="pdf")
                page_limit = min(doc.page_count, MAX_PDF_PAGES, 10)  # Limit for vision API cost
                
                # Convert pages to images for Claude Vision
                page_images = []
                for i in range(page_limit):
                    page = doc.load_page(i)
                    # Render page as PNG (300 DPI for good quality)
                    pix = page.get_pixmap(matrix=fitz.Matrix(300/72, 300/72))
                    page_images.append(pix.tobytes("png"))
                
                claude_text = await extract_with_claude_vision(page_images)
                if claude_text and len(claude_text.strip()) >= 50:
                    return file_ext, claude_text
            except Exception as claude_error:
                print(f"Claude Vision fallback failed: {claude_error}")
                # Continue to next fallback

        # Try PyPDF2 first, but fall back to pdfplumber on ANY failure (PyPDF2 can throw UnicodeDecodeError on some PDFs)
        py_pdf2_error = None
        try:
            import PyPDF2
            from io import BytesIO
            
            pdf_file = BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text_content = ""
            
            page_limit = min(len(pdf_reader.pages), MAX_PDF_PAGES)
            for page_num in range(page_limit):
                page = pdf_reader.pages[page_num]
                text_content += f"\n--- Page {page_num + 1} ---\n"
                extracted = page.extract_text()
                if extracted:
                    text_content += extracted
            
            if not text_content.strip():
                raise HTTPException(status_code=400, detail="PDF appears to be empty or image-based. Could not extract text.")
        except Exception as e:
            py_pdf2_error = e
            # Fallback: try pdfplumber
            try:
                import pdfplumber
                from io import BytesIO
                
                pdf_file = BytesIO(content)
                text_content = ""
                
                with pdfplumber.open(pdf_file) as pdf:
                    page_limit = min(len(pdf.pages), MAX_PDF_PAGES)
                    for page_num in range(page_limit):
                        page = pdf.pages[page_num]
                        text_content += f"\n--- Page {page_num + 1} ---\n"
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text
                
                if not text_content.strip():
                    raise HTTPException(status_code=400, detail="PDF appears to be empty or image-based. Could not extract text.")
            except ImportError:
                raise HTTPException(
                    status_code=500,
                    detail="PDF support requires PyPDF2 or pdfplumber. Install with: pip install PyPDF2 pdfplumber"
                )
            except Exception as plumber_error:
                # If text extraction failed, try OCR (scanned/image PDFs)
                ocr_text = try_ocr_pdf_bytes(content)
                if ocr_text and len(ocr_text.strip()) >= 50:
                    return file_ext, ocr_text

                # If OCR didn't help, surface a helpful error
                raise HTTPException(
                    status_code=400,
                    detail=(
                        "Could not extract text from PDF (likely scanned/image-only or corrupted), and OCR also failed or returned empty.\n"
                        f"PyMuPDF error: {str(pymupdf_error)}; PyPDF2 error: {str(py_pdf2_error)}; pdfplumber error: {str(plumber_error)}\n"
                        "Fix: upload the original XLSX/CSV, or OCR/export a searchable PDF."
                    )
                )
    elif file_ext in ['csv', 'txt', 'json']:
        # Regular text files (CSV, TXT, JSON)
        try:
            text_content = content.decode('utf-8')
        except UnicodeDecodeError:
            # Try other encodings
            try:
                text_content = content.decode('latin-1')
            except:
                raise HTTPException(status_code=400, detail="Could not decode file. Please ensure it's a text-based file (CSV, TXT, JSON).")
    else:
        # Unsupported format - explicitly prevent binary files from being decoded as text
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported file format: {file_ext or 'unknown'}. Supported formats: CSV, TXT, JSON, PDF, XLSX, XLS. If you uploaded an Excel file, make sure openpyxl (for XLSX) or xlrd (for XLS) is installed: pip install openpyxl xlrd"
        )
    
    # Ensure text_content is set before creating request
    if text_content is None:
        raise HTTPException(status_code=500, detail="Internal error: text_content was not set during file processing.")
    
    return file_ext, text_content

@app.get("/")
async def root():
    return {
        "message": "Ollama Data Converter API",
        "version": "1.0.0",
        "endpoints": {
            "/convert": "POST - Convert unstructured data",
            "/health": "GET - Health check",
            "/models": "GET - List available Ollama models"
        }
    }

@app.get("/health")
async def health_check():
    """Check which converter provider is available"""
    if ANTHROPIC_API_KEY:
        return {
            "status": "healthy",
            "available": True,
            "provider": "claude",
            "models": [ANTHROPIC_MODEL],
            "error": None,
        }
    if CONVERTER_PROVIDER == "claude":
        return {
            "status": "unhealthy",
            "available": False,
            "provider": "claude",
            "models": [ANTHROPIC_MODEL],
            "error": "ANTHROPIC_API_KEY not set",
        }

    try:
        model_list = await fetch_ollama_model_names()
        return {
            "status": "healthy",
            "available": True,
            "provider": "ollama",
            "models": model_list,
            "ollama_host": OLLAMA_HOST,
            "preferred_model": PREFERRED_OLLAMA_MODEL,
        }
    except Exception as e:
        return {
            "status": "unhealthy",
            "available": False,
            "provider": "ollama",
            "error": str(e),
            "ollama_host": OLLAMA_HOST,
            "preferred_model": PREFERRED_OLLAMA_MODEL,
        }

@app.get("/models")
async def list_models():
    """List available Ollama models"""
    try:
        model_names = await fetch_ollama_model_names()
        return {
            "models": [{"name": n} for n in model_names],
            "ollama_host": OLLAMA_HOST,
            "preferred_model": PREFERRED_OLLAMA_MODEL,
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to list models: {str(e)}")

def try_direct_csv_parse(text_data: str, data_type: Optional[str]) -> Optional[ConversionResponse]:
    """
    Try to parse CSV directly without Ollama if headers are clear.
    Returns ConversionResponse if successful, None if uncertain.
    """
    def normalize_header(value: str) -> str:
        if not value:
            return ""
        lowered = value.lower()
        lowered = re.sub(r'\[.*?\]', ' ', lowered)
        lowered = re.sub(r'[^a-z0-9\s]', ' ', lowered)
        lowered = re.sub(r'\s+', ' ', lowered).strip()
        return lowered

    try:
        # Parse rows with csv.reader to detect header row even if file starts with blank lines
        raw_rows = list(csv.reader(StringIO(text_data)))
        if not raw_rows:
            return None

        header_idx = None
        normalized_rows = []
        for idx, row in enumerate(raw_rows):
            normalized = [normalize_header(cell) for cell in row]
            normalized_rows.append(normalized)
            if not any(normalized):
                continue
            # Look for known header signals
            if (
                ("investor name" in normalized or "firm name" in normalized) and
                any("team member" in h or "member" == h for h in normalized)
            ):
                header_idx = idx
                break
            if ("company name" in normalized and any("funding" in h or "stage" == h for h in normalized)):
                header_idx = idx
                break
            if ("full name" in normalized and "email" in normalized):
                header_idx = idx
                break
            if (("contact name" in normalized or "contact" in normalized) and ("firm name" in normalized or "company name" in normalized)):
                header_idx = idx
                break

        if header_idx is None:
            # Fallback to DictReader using first line as headers
            reader = csv.DictReader(StringIO(text_data))
            rows = list(reader)
        else:
            header = raw_rows[header_idx]
            data_rows = raw_rows[header_idx + 1 :]
            rows = []
            for row in data_rows:
                if not any(cell.strip() for cell in row if isinstance(cell, str)) and not any(row):
                    continue
                record = {header[i]: (row[i] if i < len(row) else "") for i in range(len(header))}
                rows.append(record)

        if not rows:
            return None

        # Check first row to determine type
        first_row = rows[0]
        headers_lower = {normalize_header(k): k for k in first_row.keys() if k}
        
        print(f"[DEBUG] CSV Headers detected: {list(headers_lower.keys())}")
        
        # Detect type based on headers
        has_mentor_headers = any(h in headers_lower for h in ['full name', 'fullname']) and any(h in headers_lower for h in ['email'])
        has_corporate_headers = any(h in headers_lower for h in ['contact name', 'contactname']) and any(h in headers_lower for h in ['firm name', 'firmname', 'company name', 'companyname'])
        has_investor_headers = any(h in headers_lower for h in ['investor name', 'firm name', 'firmname']) and any(h in headers_lower for h in ['member name', 'membername', 'team member'])
        has_startup_headers = any(h in headers_lower for h in ['company name', 'companyname']) and any(h in headers_lower for h in ['funding', 'stage'])
        
        print(f"[DEBUG] Mentor headers: {has_mentor_headers}, Corporate: {has_corporate_headers}, Investor: {has_investor_headers}, Startup: {has_startup_headers}")
        
        startups = []
        investors = []
        mentors = []
        corporates = []
        warnings = []
        
        if has_mentor_headers:
            for row in rows:
                try:
                    mentor = normalize_mentor_data(row)
                    if mentor.fullName and mentor.email:
                        mentors.append(mentor)
                except Exception as e:
                    warnings.append(f"Error parsing mentor row: {str(e)}")
            
            if mentors:
                return ConversionResponse(
                    startups=[],
                    investors=[],
                    mentors=mentors,
                    corporates=[],
                    detectedType="mentor",
                    confidence=0.95,
                    warnings=warnings,
                    errors=[]
                )
        
        elif has_corporate_headers:
            for row in rows:
                try:
                    corp = normalize_corporate_data(row)
                    if corp.firmName and corp.contactName:
                        corporates.append(corp)
                except Exception as e:
                    warnings.append(f"Error parsing corporate row: {str(e)}")
            
            if corporates:
                return ConversionResponse(
                    startups=[],
                    investors=[],
                    mentors=[],
                    corporates=corporates,
                    detectedType="corporate",
                    confidence=0.95,
                    warnings=warnings,
                    errors=[]
                )
        
        elif has_investor_headers:
            for row in rows:
                try:
                    inv = normalize_investor_data(row)
                    if inv.firmName:
                        if not inv.memberName:
                            inv.memberName = "UNKNOWN"
                        investors.append(inv)
                except Exception as e:
                    warnings.append(f"Error parsing investor row: {str(e)}")
            
            if investors:
                return ConversionResponse(
                    startups=[],
                    investors=investors,
                    mentors=[],
                    corporates=[],
                    detectedType="investor",
                    confidence=0.95,
                    warnings=warnings,
                    errors=[]
                )
        
        elif has_startup_headers:
            for row in rows:
                try:
                    startup = normalize_startup_data(row)
                    if startup.companyName:
                        startups.append(startup)
                except Exception as e:
                    warnings.append(f"Error parsing startup row: {str(e)}")
            
            if startups:
                return ConversionResponse(
                    startups=startups,
                    investors=[],
                    mentors=[],
                    corporates=[],
                    detectedType="startup",
                    confidence=0.95,
                    warnings=warnings,
                    errors=[]
                )
        
        # If we couldn't determine type or no data, return None to fall back to Ollama
        return None
        
    except Exception as e:
        print(f"Direct CSV parse exception: {e}")
        return None

@app.post("/convert", response_model=ConversionResponse)
async def convert_data(request: ConversionRequest):
    """
    Convert unstructured data to structured format using Ollama
    """
    # Try direct CSV parsing first if format is CSV
    if request.format == 'csv':
        try:
            direct_result = try_direct_csv_parse(request.data, request.dataType)
            if direct_result:
                return direct_result
        except Exception as e:
            # If direct CSV parsing fails, fall through to Ollama
            print(f"Direct CSV parse failed, falling back to Ollama: {e}")
    
    try:
        # Create prompt
        prompt = create_conversion_prompt(request.data, request.dataType)

        # Decide which provider to use
        use_claude = ANTHROPIC_API_KEY is not None and ANTHROPIC_API_KEY.strip() != ""
        if CONVERTER_PROVIDER == "claude" or use_claude:
            response_text = await call_anthropic(prompt)
            try:
                parsed_data = parse_ollama_response(response_text)
            except Exception:
                retry_prompt = (
                    "Return ONLY valid JSON. Do not include markdown or explanations. "
                    "Restart the JSON from scratch and ensure all brackets are closed.\n\n"
                    + create_conversion_prompt(request.data, request.dataType)
                )
                retry_text = await call_anthropic(retry_prompt)
                parsed_data = parse_ollama_response(retry_text)
        else:
            # Check models via HTTP API (more reliable than python ollama.list on some setups)
            try:
                available_models = await fetch_ollama_model_names()
            except Exception as e:
                raise HTTPException(
                    status_code=503,
                    detail=f"Ollama not reachable at {OLLAMA_HOST}. Error: {str(e)}"
                )

            if not available_models:
                # Final fallback: attempt python client list directly
                try:
                    client = get_ollama_client()
                    models = client.list()
                    if isinstance(models, dict):
                        for m in models.get("models", []) or []:
                            name = None
                            if isinstance(m, dict) and m.get("name"):
                                name = m["name"]
                            elif isinstance(m, str):
                                name = m
                            if name:
                                available_models.append(name)
                except Exception:
                    pass

            if not available_models:
                raise HTTPException(
                    status_code=503,
                    detail=f"No Ollama models available at {OLLAMA_HOST}. Run: ollama pull llama3.1"
                )

            model_name = pick_model(available_models)

            # Call Ollama
            client = get_ollama_client()
            # Prefer JSON mode if supported by the client/version, otherwise fall back.
            chat_kwargs = dict(
                model=model_name,
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user", "content": prompt},
                ],
                options={
                    "temperature": 0.1,  # Low temperature for consistent extraction
                    "num_predict": 4096,  # More headroom to avoid truncated JSON
                },
            )
            try:
                response = client.chat(**chat_kwargs, format="json")
            except TypeError:
                response = client.chat(**chat_kwargs)

            # Extract response content
            response_text = response.get('message', {}).get('content')
            if not isinstance(response_text, str):
                raise HTTPException(status_code=502, detail="Ollama returned empty content. Ensure the model is available and retry.")

            # Parse JSON from response (retry once if model output is truncated/non-JSON)
            try:
                parsed_data = parse_ollama_response(response_text)
            except Exception:
                # Retry with a stricter prompt and higher output budget
                retry_prompt = (
                    "Return ONLY valid JSON. Do not include markdown or explanations. "
                    "Restart the JSON from scratch and ensure all brackets are closed.\n\n"
                    + create_conversion_prompt(request.data, request.dataType)
                )
                retry_kwargs = dict(
                    model=model_name,
                    messages=[
                        {"role": "system", "content": SYSTEM_PROMPT},
                        {"role": "user", "content": retry_prompt},
                    ],
                    options={
                        "temperature": 0.0,
                        "num_predict": 8192,
                    },
                )
                try:
                    retry_res = client.chat(**retry_kwargs, format="json")
                except TypeError:
                    retry_res = client.chat(**retry_kwargs)
                retry_text = retry_res.get("message", {}).get("content")
                if not isinstance(retry_text, str):
                    raise HTTPException(status_code=502, detail="Ollama returned empty content on retry.")
                parsed_data = parse_ollama_response(retry_text)
        
        # If the model returns a wrapper object (common), unwrap it.
        # Supported shapes:
        # - { startups: [...], investors: [...] }
        # - { data: [...] }
        # - { detectedType: "...", investors: [...] } etc.
        if isinstance(parsed_data, dict):
            wrapper = parsed_data

            # Direct "data" wrapper
            if isinstance(wrapper.get("data"), list):
                parsed_data = wrapper["data"]
            # Direct "startups"/"investors" wrapper: bypass generic detection loop
            elif isinstance(wrapper.get("startups"), (list, dict)) or isinstance(wrapper.get("investors"), (list, dict)):
                startups = []
                investors = []
                mentors = []
                corporates = []
                warnings = []
                errors = []

                def ensure_list(v: Any) -> List[Any]:
                    if v is None:
                        return []
                    if isinstance(v, list):
                        return v
                    return [v]

                for item in ensure_list(wrapper.get("startups")):
                    if isinstance(item, dict):
                        try:
                            s = normalize_startup_data(item)
                            if s.companyName:
                                startups.append(s)
                        except Exception as e:
                            errors.append(f"Error processing startup item: {str(e)}")

                for item in ensure_list(wrapper.get("investors")):
                    if isinstance(item, dict):
                        try:
                            inv = normalize_investor_data(item)
                            if inv.firmName:
                                if not inv.memberName:
                                    inv.memberName = "UNKNOWN"
                                    warnings.append(
                                        f"Investor missing memberName; using placeholder 'UNKNOWN' for firm '{inv.firmName}'."
                                    )
                                investors.append(inv)
                        except Exception as e:
                            errors.append(f"Error processing investor item: {str(e)}")

                for item in ensure_list(wrapper.get("mentors")):
                    if isinstance(item, dict):
                        try:
                            mentor = normalize_mentor_data(item)
                            if mentor.fullName and mentor.email:
                                mentors.append(mentor)
                        except Exception as e:
                            errors.append(f"Error processing mentor item: {str(e)}")

                for item in ensure_list(wrapper.get("corporates")):
                    if isinstance(item, dict):
                        try:
                            corp = normalize_corporate_data(item)
                            if corp.firmName and corp.contactName:
                                corporates.append(corp)
                        except Exception as e:
                            errors.append(f"Error processing corporate item: {str(e)}")

                # Determine detected type
                types_found = []
                if startups:
                    types_found.append("startup")
                if investors:
                    types_found.append("investor")
                if mentors:
                    types_found.append("mentor")
                if corporates:
                    types_found.append("corporate")
                
                detected_type = "+".join(types_found) if types_found else "unknown"
                
                if not (startups or investors or mentors or corporates):
                    errors.append("No valid data extracted. Please check the input format and column names.")

                return ConversionResponse(
                    startups=startups,
                    investors=investors,
                    mentors=mentors,
                    corporates=corporates,
                    detectedType=detected_type,
                    confidence=0.8 if (startups or investors) else 0.0,
                    warnings=warnings,
                    errors=errors,
                )

            # Fallback: treat wrapper as a single item
            else:
                parsed_data = [wrapper]
        # Normalize to list if single object
        elif isinstance(parsed_data, dict):
            parsed_data = [parsed_data]
        
        # Convert to structured format
        startups = []
        investors = []
        mentors = []
        corporates = []
        warnings = []
        errors = []
        detected_type = request.dataType or "unknown"
        
        for item in parsed_data:
            # Skip non-dict items to avoid type errors
            if not isinstance(item, dict):
                warnings.append(f"Skipping non-dict item: {item}")
                continue
            try:
                # Auto-detect type if not specified
                if not request.dataType:
                    has_startup_fields = any(k in item for k in ['companyName', 'fundingTarget', 'fundingStage'])
                    has_investor_fields = any(k in item for k in ['firmName', 'minTicketSize', 'maxTicketSize', 'memberName'])
                    has_mentor_fields = any(k in item for k in ['fullName', 'Full Name', 'expertiseAreas', 'Expertise Areas']) and any(k in item for k in ['email', 'Email'])
                    has_corporate_fields = any(k in item for k in ['contactName', 'Contact Name', 'partnershipTypes', 'Partnership Types'])
                    
                    if has_mentor_fields:
                        detected_type = "mentor"
                    elif has_corporate_fields:
                        detected_type = "corporate"
                    elif has_startup_fields and not has_investor_fields:
                        detected_type = "startup"
                    elif has_investor_fields and not has_startup_fields:
                        detected_type = "investor"
                    elif has_startup_fields and has_investor_fields:
                        # Ambiguous - check more indicators
                        if 'companyName' in item and 'fundingTarget' in item:
                            detected_type = "startup"
                        else:
                            detected_type = "investor"
                
                # Convert based on detected type
                if detected_type == "startup" or (not request.dataType and 'companyName' in item):
                    startup = normalize_startup_data(item)
                    if startup.companyName:
                        startups.append(startup)
                elif detected_type == "mentor" or (not request.dataType and any(k in item for k in ['fullName', 'Full Name'])):
                    mentor = normalize_mentor_data(item)
                    if mentor.fullName and mentor.email:
                        mentors.append(mentor)
                    elif mentor.fullName:
                        warnings.append(f"Mentor '{mentor.fullName}' missing email, skipping")
                elif detected_type == "corporate" or (not request.dataType and any(k in item for k in ['contactName', 'Contact Name'])):
                    corporate = normalize_corporate_data(item)
                    if corporate.firmName and corporate.contactName:
                        corporates.append(corporate)
                    elif corporate.firmName:
                        warnings.append(f"Corporate '{corporate.firmName}' missing contact name, skipping")
                elif detected_type == "investor" or (not request.dataType and 'firmName' in item):
                    investor = normalize_investor_data(item)
                    if investor.firmName:
                        # Some sources (esp. PDFs) list only firm names without a specific person.
                        # Don't hard-fail the whole conversion; fill a placeholder and warn.
                        if not investor.memberName:
                            investor.memberName = "UNKNOWN"
                            warnings.append(
                                f"Investor missing memberName; using placeholder 'UNKNOWN' for firm '{investor.firmName}'."
                            )
                        investors.append(investor)
                else:
                    warnings.append(f"Could not determine type for item: {item}")
            except Exception as e:
                errors.append(f"Error processing item: {str(e)}")
        
        if not (startups or investors or mentors or corporates):
            errors.append("No valid data extracted. Please check the input format and column names.")
        
        return ConversionResponse(
            startups=startups,
            investors=investors,
            mentors=mentors,
            corporates=corporates,
            detectedType=detected_type,
            confidence=0.8 if (startups or investors) else 0.0,
            warnings=warnings,
            errors=errors
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")

def validate_structured_rows(
    startups: List[StartupData], 
    investors: List[InvestorData],
    mentors: List[MentorData] = None,
    corporates: List[CorporateData] = None
) -> List[str]:
    """
    Row-level validation with explicit row numbers and missing fields.
    """
    errors: List[str] = []

    for idx, s in enumerate(startups, start=1):
        missing = []
        if not s.companyName:
            missing.append("companyName")
        # NOTE: We intentionally do NOT hard-require every field here.
        # Many PDFs / unstructured sources omit fields like stage/geo/ticket size.
        # The UI supports editing later; blocking imports is worse UX.
        if missing:
            errors.append(f"Startup row {idx}: missing {', '.join(missing)}")

    for idx, inv in enumerate(investors, start=1):
        missing = []
        if not inv.firmName:
            missing.append("firmName")
        if not inv.memberName:
            missing.append("memberName")
        # Do not require geoFocus/industryPreferences/stagePreferences/ticket sizes here.
        if missing:
            errors.append(f"Investor row {idx}: missing {', '.join(missing)}")

    for idx, mentor in enumerate(mentors or [], start=1):
        missing = []
        if not mentor.fullName:
            missing.append("fullName")
        if not mentor.email:
            missing.append("email")
        if missing:
            errors.append(f"Mentor row {idx}: missing {', '.join(missing)}")

    for idx, corp in enumerate(corporates or [], start=1):
        missing = []
        if not corp.firmName:
            missing.append("firmName")
        if not corp.contactName:
            missing.append("contactName")
        if missing:
            errors.append(f"Corporate row {idx}: missing {', '.join(missing)}")

    return errors

def build_startup_csv(startups: List[StartupData]) -> str:
    headers = ["company_name", "geo_markets", "industry", "funding_target", "funding_stage"]
    output = StringIO()
    writer = csv.DictWriter(output, fieldnames=headers)
    writer.writeheader()
    for s in startups:
        writer.writerow({
            "company_name": s.companyName or "",
            "geo_markets": "; ".join(s.geoMarkets) if s.geoMarkets else "",
            "industry": s.industry or "",
            "funding_target": s.fundingTarget if s.fundingTarget is not None else "",
            "funding_stage": s.fundingStage or "",
        })
    return output.getvalue()

def build_investor_csv(investors: List[InvestorData]) -> str:
    headers = [
        "firm_name",
        "investment_member",
        "geo_focus",
        "industry_preferences",
        "min_ticket_size",
        "max_ticket_size",
        "total_slots",
        "table_number",
    ]
    output = StringIO()
    writer = csv.DictWriter(output, fieldnames=headers)
    writer.writeheader()
    for inv in investors:
        writer.writerow({
            "firm_name": inv.firmName or "",
            "investment_member": inv.memberName or "",
            "geo_focus": "; ".join(inv.geoFocus) if inv.geoFocus else "",
            "industry_preferences": "; ".join(inv.industryPreferences) if inv.industryPreferences else "",
            "min_ticket_size": inv.minTicketSize if inv.minTicketSize is not None else "",
            "max_ticket_size": inv.maxTicketSize if inv.maxTicketSize is not None else "",
            "total_slots": inv.totalSlots if inv.totalSlots is not None else "",
            "table_number": inv.tableNumber or "",
        })
    return output.getvalue()

@app.post("/convert-file")
async def convert_file(file: UploadFile = File(...), dataType: Optional[str] = None):
    """Convert uploaded file (CSV, text, PDF, etc.)"""
    try:
        file_ext, text_content = await extract_text_content(file)
        request = ConversionRequest(
            data=text_content,
            dataType=dataType,
            format=file_ext
        )
        conversion_result = await convert_data(request)
        # Include extracted text for downstream indexing (truncate to control payload size)
        conversion_result.raw_content = text_content[:MAX_MODEL_INPUT_CHARS]

        # Validate critical identifiers, but don't hard-fail if optional fields are missing.
        row_errors = validate_structured_rows(
            conversion_result.startups, 
            conversion_result.investors,
            conversion_result.mentors,
            conversion_result.corporates
        )

        # Block only if nothing was extracted
        has_any_data = (
            conversion_result.startups or 
            conversion_result.investors or 
            conversion_result.mentors or 
            conversion_result.corporates
        )
        if not has_any_data:
            # Return a 200 with errors so the frontend can show a meaningful message
            # (instead of a generic HTTP failure that hides conversion_result.errors).
            conversion_result.errors = (conversion_result.errors or []) + [
                "No valid data extracted. Please check the input format and column names. Expected columns for Investors: firmName, memberName, geoFocus, industryPreferences, stagePreferences. For Mentors: fullName, email, geoFocus. For Corporates: firmName, contactName, geoFocus, partnershipTypes."
            ]
            return conversion_result

        # Surface missing critical fields as warnings so users can import and edit in the UI.
        if row_errors:
            conversion_result.warnings = (conversion_result.warnings or []) + row_errors

        return conversion_result
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"File conversion failed: {str(e)}")

def parse_google_drive_url(url: str) -> Tuple[str, str]:
    patterns = [
        ("document", r"https?://docs\.google\.com/document/d/([^/]+)"),
        ("presentation", r"https?://docs\.google\.com/presentation/d/([^/]+)"),
        ("spreadsheet", r"https?://docs\.google\.com/spreadsheets/d/([^/]+)"),
        ("drive", r"https?://drive\.google\.com/file/d/([^/]+)"),
    ]
    for kind, pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return kind, match.group(1)
    # Alternate Drive URL pattern: open?id=FILE_ID
    match = re.search(r"[?&]id=([^&]+)", url)
    if match:
        return "drive", match.group(1)
    raise HTTPException(status_code=400, detail="Unsupported Google Drive URL format.")

@app.post("/ingest/clickup", response_model=ClickUpIngestResponse)
async def ingest_clickup(request: ClickUpIngestRequest):
    if not CLICKUP_API_TOKEN:
        raise HTTPException(status_code=503, detail="CLICKUP_API_TOKEN not set on the server.")

    list_id = request.list_id.strip()
    if not list_id:
        raise HTTPException(status_code=400, detail="list_id is required.")

    url = f"https://api.clickup.com/api/v2/list/{list_id}/task"
    params = {"include_closed": "true" if request.include_closed else "false"}
    headers = {"Authorization": CLICKUP_API_TOKEN}

    async with httpx.AsyncClient(timeout=30.0) as client:
        res = await client.get(url, headers=headers, params=params)
        if res.status_code >= 400:
            raise HTTPException(status_code=res.status_code, detail=res.text[:400])
        data = res.json()

    tasks = []
    for task in data.get("tasks", []):
        assignees = [a.get("username") for a in (task.get("assignees") or []) if a.get("username")]
        tasks.append({
            "id": task.get("id"),
            "name": task.get("name"),
            "url": task.get("url"),
            "status": (task.get("status") or {}).get("status"),
            "assignees": assignees,
            "description": task.get("description"),
        })

    return ClickUpIngestResponse(tasks=tasks)

@app.post("/ingest/clickup/lists", response_model=ClickUpListsResponse)
async def list_clickup_lists(request: ClickUpListsRequest):
    if not CLICKUP_API_TOKEN:
        raise HTTPException(status_code=503, detail="CLICKUP_API_TOKEN not set on the server.")

    team_id = request.team_id.strip()
    if not team_id:
        raise HTTPException(status_code=400, detail="team_id is required.")

    url = f"https://api.clickup.com/api/v2/team/{team_id}/list"
    headers = {"Authorization": CLICKUP_API_TOKEN}

    async with httpx.AsyncClient(timeout=30.0) as client:
        res = await client.get(url, headers=headers, params={"archived": "false"})
        if res.status_code >= 400:
            raise HTTPException(status_code=res.status_code, detail=res.text[:400])
        data = res.json()

    lists = []
    for item in data.get("lists", []):
        lists.append({
            "id": item.get("id"),
            "name": item.get("name"),
        })

    return ClickUpListsResponse(lists=lists)

async def stream_anthropic_answer(prompt: str, question: str = "", sources: List[AskSource] = None) -> AsyncGenerator[str, None]:
    """
    Stream Claude's response token by token for ChatGPT-like experience.
    """
    if not ANTHROPIC_API_KEY:
        yield json.dumps({"error": "ANTHROPIC_API_KEY not set"})
        return

    headers = {
        "Content-Type": "application/json",
        "x-api-key": ANTHROPIC_API_KEY,
        "anthropic-version": "2023-06-01",
    }
    url = get_anthropic_api_url()
    default_url = "https://api.anthropic.com/v1/messages"

    # Choose model based on question complexity
    is_comprehensive = is_comprehensive_question(question)
    use_haiku = question and sources and is_simple_question(question, sources) and not is_comprehensive
    model_list = ([HAIKU_MODEL] + ANTHROPIC_MODEL_FALLBACKS) if use_haiku else ANTHROPIC_MODEL_FALLBACKS
    if is_comprehensive:
        max_tokens = 8000  # Comprehensive questions need space for detailed memos/analysis
    else:
        max_tokens = 2000 if use_haiku else ASK_MAX_TOKENS  # Haiku also gets more room
    
    async with httpx.AsyncClient(timeout=60.0) as client:
        for model_name in [m for m in model_list if m]:
            payload = {
                "model": model_name,
                "max_tokens": max_tokens,
                "temperature": 0.1,
                "stream": True,  # Enable streaming
                "system": "You are Orbit AI, a VC intelligence system. You answer questions STRICTLY from provided sources only. Never use general knowledge. If information isn't in the sources, say so explicitly. Always cite sources with [1], [2], etc.",
                "messages": [
                    {"role": "user", "content": prompt}
                ],
            }

            try:
                async with client.stream("POST", url, headers=headers, json=payload) as response:
                    if response.status_code == 404 and url != default_url:
                        async with client.stream("POST", default_url, headers=headers, json=payload) as retry_response:
                            async for line in retry_response.aiter_lines():
                                if line.startswith("data: "):
                                    data_str = line[6:]
                                    if data_str == "[DONE]":
                                        return
                                    try:
                                        data = json.loads(data_str)
                                        if "delta" in data and "text" in data["delta"]:
                                            yield json.dumps({"text": data["delta"]["text"]})
                                    except json.JSONDecodeError:
                                        continue
                        return
                    
                    if response.status_code >= 400:
                        error_text = await response.aread()
                        error_str = error_text[:200].decode() if isinstance(error_text, bytes) else str(error_text)[:200]
                        # If it's a 404 for model not found, try next model instead of failing
                        if response.status_code == 404 and ("not_found_error" in error_str.lower() or "model:" in error_str.lower()):
                            if model_name == model_list[-1]:  # Last model
                                yield json.dumps({"error": f"Claude API error ({response.status_code}): All models failed. Please check your API access. Error: {error_str}"})
                                return
                            # Skip this invalid model and try next one
                            continue
                        yield json.dumps({"error": f"Claude API error ({response.status_code}): {error_str}"})
                        return

                    async for line in response.aiter_lines():
                        if line.startswith("data: "):
                            data_str = line[6:]
                            if data_str == "[DONE]":
                                return
                            try:
                                data = json.loads(data_str)
                                if "delta" in data and "text" in data["delta"]:
                                    yield json.dumps({"text": data["delta"]["text"]})
                                elif "error" in data:
                                    yield json.dumps({"error": str(data["error"])})
                                    return
                            except json.JSONDecodeError:
                                continue
                    return  # Success
            except Exception as e:
                if model_name == model_list[-1]:  # Last model, yield error
                    yield json.dumps({"error": f"All models failed: {str(e)}"})
                continue


@app.post("/ask", response_model=AskResponse)
async def ask_fund(request: AskRequest):
    question = (request.question or "").strip()
    if not question:
        raise HTTPException(status_code=400, detail="question is required.")

    no_info_message = "I don't have information about this in the provided sources. Please upload relevant documents or try a different question."
    # Use LLM-based query rewriting for better pronoun resolution
    resolved_question = await rewrite_query_with_llm(question, request.previous_messages or [])
    if not is_meta_question(resolved_question) and not has_question_overlap(
        resolved_question, request.sources or [], request.previous_messages or [], request.decisions or []
    ):
        return AskResponse(answer=no_info_message)

    prompt = build_answer_prompt(resolved_question, request.sources, request.decisions, request.previous_messages)
    answer = await call_anthropic_answer(prompt, question=resolved_question, sources=request.sources)
    return AskResponse(answer=answer)


@app.post("/ask/stream")
async def ask_fund_stream(request: AskRequest):
    """
    Streaming endpoint for ChatGPT-like gradual text typing.
    Returns Server-Sent Events (SSE) stream.
    """
    try:
        question = (request.question or "").strip()
        if not question:
            raise HTTPException(status_code=400, detail="question is required.")

        no_info_message = "I don't have information about this in the provided sources. Please upload relevant documents or try a different question."
        # Use LLM-based query rewriting for better pronoun resolution
        print(f"[DEBUG] /ask/stream - Original question: {question}")
        print(f"[DEBUG] /ask/stream - Previous messages count: {len(request.previous_messages or [])}")
        
        resolved_question = await rewrite_query_with_llm(question, request.previous_messages or [])
        print(f"[DEBUG] /ask/stream - Resolved question: {resolved_question}")
        
        # Check overlap - but be more lenient for follow-up questions
        has_overlap = has_question_overlap(
            resolved_question, request.sources or [], request.previous_messages or [], request.decisions or []
        )
        print(f"[DEBUG] /ask/stream - has_overlap: {has_overlap}")
        
        if not is_meta_question(resolved_question) and not has_overlap:
            print(f"[DEBUG] /ask/stream - REJECTING: No overlap found, returning no_info_message")
            async def generate_empty():
                yield f"data: {json.dumps({'text': no_info_message})}\n\n"
                yield "data: [DONE]\n\n"
            return StreamingResponse(
                generate_empty(),
                media_type="text/event-stream",
                headers={
                    "Cache-Control": "no-cache",
                    "Connection": "keep-alive",
                    "X-Accel-Buffering": "no",
                    "Access-Control-Allow-Origin": "*",
                    "Access-Control-Allow-Methods": "POST, OPTIONS",
                    "Access-Control-Allow-Headers": "*",
                },
            )

        # Handle previous_messages safely (default to empty list if None)
        previous_messages = request.previous_messages or []
        
        # Debug logging - DETAILED
        print(f"[DEBUG] ========== /ask/stream REQUEST ==========")
        print(f"[DEBUG] Question: {question}")
        print(f"[DEBUG] Resolved question: {resolved_question}")
        print(f"[DEBUG] Previous messages count: {len(previous_messages)}")
        if previous_messages:
            print(f"[DEBUG] First message: {previous_messages[0].role}: {previous_messages[0].content[:200]}")
            print(f"[DEBUG] Last message: {previous_messages[-1].role}: {previous_messages[-1].content[:200]}")
            # Print all messages for debugging
            for i, msg in enumerate(previous_messages):
                print(f"[DEBUG] Message {i+1}: {msg.role}: {msg.content[:100]}...")
        else:
            print(f"[DEBUG] ⚠️⚠️⚠️ WARNING: NO PREVIOUS MESSAGES PROVIDED! ⚠️⚠️⚠️")
        print(f"[DEBUG] =========================================")
        
        prompt = build_answer_prompt(resolved_question, request.sources or [], request.decisions or [], previous_messages)
        
        async def generate():
            try:
                # Send an initial ping so the client doesn't think the stream is empty
                yield f"data: {json.dumps({'ping': True})}\n\n"
                async for chunk in stream_anthropic_answer(prompt, question=question, sources=request.sources or []):
                    yield f"data: {chunk}\n\n"
                yield "data: [DONE]\n\n"
            except Exception as e:
                import traceback
                error_trace = traceback.format_exc()
                print(f"Stream generation error: {error_trace}")
                error_msg = str(e)[:500]  # Limit error message length
                yield f'data: {{"error": "{error_msg}"}}\n\n'
        
        return StreamingResponse(
            generate(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "X-Accel-Buffering": "no",  # Disable nginx buffering
                "Access-Control-Allow-Origin": "*",  # Explicit CORS header for streaming
                "Access-Control-Allow-Methods": "POST, OPTIONS",
                "Access-Control-Allow-Headers": "*",
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"Stream endpoint error: {error_trace}")
        raise HTTPException(
            status_code=500,
            detail=f"Streaming error: {str(e)[:200]}"
        )

def normalize_embedding(embedding: List[float]) -> List[float]:
    """Ensure embeddings are a fixed size for pgvector."""
    if not embedding:
        return []
    if len(embedding) == EMBEDDING_DIM:
        return embedding
    if len(embedding) > EMBEDDING_DIM:
        return embedding[:EMBEDDING_DIM]
    # Pad with zeros if smaller than expected
    return embedding + [0.0] * (EMBEDDING_DIM - len(embedding))



@app.post("/rerank", response_model=RerankResponse)
async def rerank_documents(request: RerankRequest):
    """
    Optional cross-encoder reranking for hybrid search results.
    If no reranker API key is configured, returns the original order with score=0.
    """
    if not request.documents:
        return RerankResponse(results=[])

    if not COHERE_API_KEY:
        return RerankResponse(
            results=[RerankResult(id=d.id, score=0.0) for d in request.documents]
        )

    payload = {
        "model": RERANK_MODEL,
        "query": request.query,
        "documents": [d.text for d in request.documents],
        "top_n": request.top_n or len(request.documents),
    }
    headers = {
        "Authorization": f"Bearer {COHERE_API_KEY}",
        "Content-Type": "application/json",
    }

    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.post("https://api.cohere.ai/v1/rerank", json=payload, headers=headers)
        response.raise_for_status()
        data = response.json()
        results = []
        for item in data.get("results", []):
            idx = item.get("index")
            score = item.get("relevance_score", 0.0)
            if idx is None or idx >= len(request.documents):
                continue
            results.append(RerankResult(id=request.documents[idx].id, score=score))
        return RerankResponse(results=results)
    except Exception:
        # Fail closed: return original order if rerank fails
        return RerankResponse(
            results=[RerankResult(id=d.id, score=0.0) for d in request.documents]
        )
        
        if response.status_code >= 400:
            error_text = response.text[:400]
            raise HTTPException(
                status_code=502,
                detail=f"OpenAI embedding API error ({response.status_code}): {error_text}"
            )
        
        data = response.json()
        embedding_data = data.get("data", [{}])[0] if data.get("data") else {}
        embedding = embedding_data.get("embedding")
        
        if not embedding:
            raise HTTPException(status_code=502, detail="No embedding returned from OpenAI.")
        
        return normalize_embedding(embedding)


async def generate_embedding_voyage(text: str, input_type: str) -> List[float]:
    """Generate embedding using VoyageAI API."""
    if not VOYAGE_API_KEY:
        raise HTTPException(
            status_code=503,
            detail="VOYAGE_API_KEY not set. Set it in the server environment to use VoyageAI embeddings."
        )

    payload = {
        "model": VOYAGE_EMBEDDING_MODEL,
        "input": [text],
        "input_type": input_type,
    }

    async with httpx.AsyncClient(timeout=30.0) as client:
        response = await client.post(
            "https://api.voyageai.com/v1/embeddings",
            headers={
                "Authorization": f"Bearer {VOYAGE_API_KEY}",
                "Content-Type": "application/json",
            },
            json=payload,
        )

        if response.status_code >= 400:
            error_text = response.text[:400]
            raise HTTPException(
                status_code=502,
                detail=f"VoyageAI embedding API error ({response.status_code}): {error_text}"
            )

        data = response.json() or {}
        embedding_data = data.get("data", [{}])[0] if data.get("data") else {}
        embedding = embedding_data.get("embedding")

        if not embedding:
            raise HTTPException(status_code=502, detail="No embedding returned from VoyageAI.")

        return normalize_embedding(embedding)

async def generate_embedding_ollama(text: str) -> List[float]:
    """Generate embedding using Ollama."""
    try:
        response = ollama.embeddings(model=OLLAMA_EMBEDDING_MODEL, prompt=text)
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Ollama embedding failed: {str(e)}")

    embedding = response.get("embedding") if isinstance(response, dict) else None
    if not embedding:
        raise HTTPException(status_code=502, detail="No embedding returned from Ollama.")
    return normalize_embedding(embedding)

@app.post("/rewrite-query", response_model=RewriteQueryResponse)
async def rewrite_query_endpoint(request: RewriteQueryRequest):
    """
    Rewrite a query using LLM to resolve pronouns and contextualize vague questions.
    This should be called BEFORE document search to ensure proper query understanding.
    """
    question = (request.question or "").strip()
    if not question:
        raise HTTPException(status_code=400, detail="question is required.")
    
    rewritten = await rewrite_query_with_llm(question, request.previous_messages or [])
    return RewriteQueryResponse(rewritten_question=rewritten)

@app.post("/embed/query", response_model=EmbedResponse)
async def embed_query(request: EmbedRequest):
    text = (request.text or "").strip()
    if not text:
        raise HTTPException(status_code=400, detail="text is required.")

    try:
        input_type = (request.input_type or "document").strip().lower()
        if input_type not in ["document", "query"]:
            input_type = "document"

        if EMBEDDINGS_PROVIDER == "voyage":
            embedding = await generate_embedding_voyage(text, input_type)
        elif EMBEDDINGS_PROVIDER == "openai":
            embedding = await generate_embedding_openai(text)
        elif EMBEDDINGS_PROVIDER == "ollama":
            embedding = await generate_embedding_ollama(text)
        else:
            raise HTTPException(
                status_code=503,
                detail=(
                    f"EMBEDDINGS_PROVIDER '{EMBEDDINGS_PROVIDER}' not supported. "
                    "Use 'voyage', 'openai', or 'ollama'."
                )
            )
        
        return EmbedResponse(embedding=embedding)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=503,
            detail=f"Embedding generation failed: {str(e)}"
        )

@app.post("/ingest/google-drive", response_model=GoogleDriveIngestResponse)
async def ingest_google_drive(request: GoogleDriveIngestRequest):
    if not request.access_token:
        raise HTTPException(status_code=400, detail="Access token required for private Google Drive files.")
    
    kind, file_id = parse_google_drive_url(request.url.strip())
    
    if kind not in ["document", "presentation", "spreadsheet"]:
        raise HTTPException(status_code=400, detail="Drive file URLs require a Docs/Slides/Sheets link.")

    # Use Google Drive API v3 to download the file
    # For Google Docs/Sheets/Slides, we need to export them
    if kind == "document":
        # Export as plain text
        api_url = f"https://www.googleapis.com/drive/v3/files/{file_id}/export?mimeType=text/plain"
        source_type = "notes"
    elif kind == "presentation":
        # Export as plain text
        api_url = f"https://www.googleapis.com/drive/v3/files/{file_id}/export?mimeType=text/plain"
        source_type = "deck"
    elif kind == "spreadsheet":
        # Export as CSV
        api_url = f"https://www.googleapis.com/drive/v3/files/{file_id}/export?mimeType=text/csv"
        source_type = "notes"
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type.")

    headers = {
        "Authorization": f"Bearer {request.access_token}",
        "Accept": "text/plain" if kind != "spreadsheet" else "text/csv"
    }

    async with httpx.AsyncClient(timeout=30.0) as client:
        res = await client.get(api_url, headers=headers)
        if res.status_code >= 400:
            error_detail = res.text[:500] if res.text else "No error details"
            raise HTTPException(
                status_code=res.status_code,
                detail=f"Google Drive API failed (status {res.status_code}): {error_detail}. Make sure you have Drive access and the file is accessible."
            )
        content = res.text
        
        # Log if content is empty
        if not content or len(content.strip()) == 0:
            print(f"WARNING: Google Drive API returned empty content for {file_id}. Status: {res.status_code}")
            content = f"[Empty content from Google Drive file: {file_id}]"

    title = f"{kind}-{file_id[:8]}"
    return GoogleDriveIngestResponse(title=title, content=content, raw_content=content, sourceType=source_type)

class ValidationRequest(BaseModel):
    data: str
    dataType: Optional[str] = None

class ValidationResponse(BaseModel):
    isValid: bool
    missingFields: Dict[str, List[str]]  # { "startups": ["geoMarkets"], "investors": ["minTicketSize"] }
    incompleteFields: Dict[str, List[str]]  # Fields that exist but are incomplete
    suggestions: List[str]  # Suggestions for what to add
    extractedData: Dict[str, Any]  # What was successfully extracted

@app.post("/validate-file", response_model=FileValidationResponse)
async def validate_file(file: UploadFile = File(...), dataType: Optional[str] = None):
    """
    Validate an uploaded file (any supported format) and return:
    - row-level errors with explicit row numbers
    - CSV templates with extracted rows prefilled and missing columns preserved
    """
    try:
        file_ext, text_content = await extract_text_content(file)
        conversion_request = ConversionRequest(
            data=text_content,
            dataType=dataType,
            format=file_ext
        )
        conversion_result = await convert_data(conversion_request)

        row_errors = validate_structured_rows(conversion_result.startups, conversion_result.investors)
        errors = (conversion_result.errors or []) + row_errors
        warnings = conversion_result.warnings or []

        startup_csv = build_startup_csv(conversion_result.startups) if conversion_result.startups else build_startup_csv([])
        investor_csv = build_investor_csv(conversion_result.investors) if conversion_result.investors else build_investor_csv([])

        return FileValidationResponse(
            isValid=len(errors) == 0,
            errors=errors,
            warnings=warnings,
            detectedType=conversion_result.detectedType,
            startupCsvTemplate=startup_csv,
            investorCsvTemplate=investor_csv,
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"File validation failed: {str(e)}")

@app.post("/validate", response_model=ValidationResponse)
async def validate_data(request: ValidationRequest):
    """
    Validate data and identify what's missing
    This is what the investment team needs - tells them what to add!
    """
    try:
        # First, try to convert the data
        conversion_request = ConversionRequest(
            data=request.data,
            dataType=request.dataType
        )
        conversion_result = await convert_data(conversion_request)
        
        missing_fields = {"startups": [], "investors": []}
        incomplete_fields = {"startups": [], "investors": []}
        suggestions = []
        
        # Check startups
        for startup in conversion_result.startups:
            startup_missing = []
            startup_incomplete = []
            
            if not startup.companyName or (isinstance(startup.companyName, str) and startup.companyName.strip() == ""):
                startup_missing.append("companyName")
            if not startup.geoMarkets or len(startup.geoMarkets) == 0:
                startup_missing.append("geoMarkets")
                suggestions.append(f"Add geographic markets for {startup.companyName}")
            if not startup.industry or (isinstance(startup.industry, str) and startup.industry.strip() == ""):
                startup_missing.append("industry")
            if not startup.fundingTarget or startup.fundingTarget == 0:
                startup_missing.append("fundingTarget")
                suggestions.append(f"Add funding target amount for {startup.companyName}")
            if not startup.fundingStage or (isinstance(startup.fundingStage, str) and startup.fundingStage.strip() == ""):
                startup_missing.append("fundingStage")
                suggestions.append(f"Add funding stage (Pre-seed, Seed, Series A, etc.) for {startup.companyName}")
            
            if startup_missing:
                missing_fields["startups"].extend(startup_missing)
        
        # Check investors
        for investor in conversion_result.investors:
            investor_missing = []
            investor_incomplete = []
            
            if not investor.firmName or (isinstance(investor.firmName, str) and investor.firmName.strip() == ""):
                investor_missing.append("firmName")
            if not investor.memberName or (isinstance(investor.memberName, str) and investor.memberName.strip() == ""):
                investor_missing.append("memberName")
                suggestions.append(f"Add investor member name (person) for {investor.firmName or 'this investor'}")
            if not investor.geoFocus or len(investor.geoFocus) == 0:
                investor_missing.append("geoFocus")
                suggestions.append(f"Add geographic focus for {investor.firmName}")
            if not investor.industryPreferences or len(investor.industryPreferences) == 0:
                investor_missing.append("industryPreferences")
                suggestions.append(f"Add industry preferences for {investor.firmName}")
            if not investor.stagePreferences or len(investor.stagePreferences) == 0:
                investor_missing.append("stagePreferences")
                suggestions.append(f"Add stage preferences (Seed, Series A, etc.) for {investor.firmName}")
            if not investor.minTicketSize or investor.minTicketSize == 0:
                investor_missing.append("minTicketSize")
                suggestions.append(f"Add minimum ticket size for {investor.firmName}")
            if not investor.maxTicketSize or investor.maxTicketSize == 0:
                investor_missing.append("maxTicketSize")
                suggestions.append(f"Add maximum ticket size for {investor.firmName}")
            if not investor.totalSlots or investor.totalSlots == 0:
                investor_missing.append("totalSlots")
                suggestions.append(f"Add number of meeting slots for {investor.firmName}")
            
            if investor_missing:
                missing_fields["investors"].extend(investor_missing)
        
        # Remove duplicates
        missing_fields["startups"] = list(set(missing_fields["startups"]))
        missing_fields["investors"] = list(set(missing_fields["investors"]))
        
        is_valid = (
            len(missing_fields["startups"]) == 0 and
            len(missing_fields["investors"]) == 0 and
            len(conversion_result.errors) == 0
        )
        
        return ValidationResponse(
            isValid=is_valid,
            missingFields=missing_fields,
            incompleteFields=incomplete_fields,
            suggestions=suggestions,
            extractedData={
                "startups": [s.dict() for s in conversion_result.startups],
                "investors": [i.dict() for i in conversion_result.investors],
                "detectedType": conversion_result.detectedType,
                "confidence": conversion_result.confidence
            }
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Validation failed: {str(e)}")

if __name__ == "__main__":
    import os
    import uvicorn
    port = int(os.environ.get("PORT", os.environ.get("OLLAMA_CONVERTER_PORT", "8000")))
    uvicorn.run(app, host="0.0.0.0", port=port)

